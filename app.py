import re
import io
import hmac
import json
import time
import base64
import hashlib
import secrets
import requests
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone, date
from typing import Dict, List, Optional, Tuple

import streamlit as st

# Optional exporters (si no están instalados, igual corre sin exportar)
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    REPORTLAB_OK = True
except Exception:
    REPORTLAB_OK = False

try:
    from docx import Document
    from docx.shared import Inches
    PYDOCX_OK = True
except Exception:
    PYDOCX_OK = False

from google.oauth2.service_account import Credentials
from google.auth.transport.requests import AuthorizedSession

# =========================
# PAGE CONFIG
# =========================
st.set_page_config(page_title="Control Comunidades", page_icon="🛡️", layout="wide")

EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
STATUS_VALUES = ["pending", "ok", "fail"]
REPORT_STATUS = ["Draft", "Final"]
ROLES = ["admin", "supervisor", "mantenedor", "analista", "conserje", "viewer"]

# =========================
# SCHEMA (TABS + HEADERS)
# =========================
TABS_HEADERS: Dict[str, List[str]] = {
    "Users": [
        "user_id", "email", "full_name",
        "is_admin", "is_active",
        "password_hash",
        "created_at", "last_login_at"
    ],
    "Invites": [
        "email", "invite_code_hash", "expires_at", "used_at",
        "created_by", "created_at"
    ],
    "UserCommunityAccess": [
        "email", "community_id",
        "role", "is_active",
        "can_create_reports", "can_view_summary"
    ],
    "Communities": [
        "community_id", "community_name",
        "is_active",
        "cover_file_id", "cover_web_view_link",
        "created_at"
    ],
    "Installations": [
        "installation_id", "community_id",
        "category", "installation_name",
        "is_active", "created_at"
    ],
    "Tasks": [
        "task_id", "community_id", "installation_id",
        "task_name", "task_order",
        "is_required",
        "is_active", "created_at"
    ],
    "Reports": [
        "report_id", "community_id", "community_name",
        "report_date",
        "status",
        "created_by",
        "created_at", "updated_at"
    ],
    "ReportItems": [
        "report_item_id",
        "report_id", "community_id",
        "installation_id", "installation_name", "category",
        "task_id", "task_name", "task_order",
        "status", "note",
        "photo_file_id", "photo_web_view_link",
        "updated_at", "updated_by"
    ],
}

DEFAULT_CATEGORIES = ["Críticos", "Accesos", "Higiene", "Comunes", "Infra"]

DEFAULT_INSTALLATIONS_AND_TASKS = [
    # (category, installation, [(task_name, order, required)])
    ("Críticos", "Sala de Bombas", [("Presión y alternancia", 10, True), ("Fugas / Sellos", 20, True)]),
    ("Críticos", "Sala de Calderas", [("Temperatura", 10, True), ("Fugas", 20, True)]),
    ("Críticos", "Generador", [("Nivel petróleo", 10, True), ("Batería", 20, True)]),
    ("Críticos", "PEAS (Presurización)", [("Prueba de ventilador", 10, True)]),
    ("Críticos", "Ascensores", [("Nivelación", 10, True), ("Limpieza rieles", 20, False)]),
    ("Accesos", "Portones", [("Sensores", 10, True), ("Velocidad", 20, True)]),
    ("Accesos", "Control Biométrico", [("Lectores huella/tarjeta", 10, True)]),
    ("Higiene", "Sala de Basura", [("Desinfección", 10, True), ("Contenedores", 20, True)]),
    ("Higiene", "Ductos", [("Cierre escotillas", 10, True)]),
    ("Comunes", "Piscina", [("Parámetros Cl/pH", 10, True)]),
    ("Comunes", "Quincho / Eventos", [("Mobiliario", 10, True), ("Higiene", 20, True)]),
    ("Comunes", "Gym / Sauna", [("Máquinas", 10, True), ("Tableros", 20, False)]),
    ("Infra", "Pasillos", [("Luces de emergencia", 10, True)]),
    ("Infra", "Subterráneo", [("Filtraciones", 10, True), ("Limpieza", 20, True)]),
    ("Infra", "Jardines", [("Riego programado", 10, True)]),
]

# =========================
# SECRETS
# =========================
SHEET_ID = st.secrets["SHEET_ID"]

APP_PEPPER = str(st.secrets.get("APP_PEPPER", "CHANGE_ME"))
INVITE_EXPIRY_HOURS = int(str(st.secrets.get("INVITE_EXPIRY_HOURS", "48")))

BOOTSTRAP_ADMIN_EMAILS = [
    x.strip().lower()
    for x in str(st.secrets.get("BOOTSTRAP_ADMIN_EMAILS", "")).split(",")
    if x.strip()
]

BOOTSTRAP_SETUP_TOKEN = str(st.secrets.get("BOOTSTRAP_SETUP_TOKEN", ""))

FOTOS_ROOT_FOLDER_ID = str(st.secrets.get("FOTOS_ROOT_FOLDER_ID", "")).strip()

# =========================
# TIME HELPERS
# =========================
def now_utc() -> datetime:
    return datetime.now(timezone.utc)

def iso(dt: datetime) -> str:
    return dt.astimezone(timezone.utc).isoformat()

def parse_iso(s: str) -> Optional[datetime]:
    if not s:
        return None
    try:
        dt = datetime.fromisoformat(s)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(timezone.utc)
    except Exception:
        return None

def norm_email(s: str) -> str:
    return (s or "").strip().lower()

def safe_bool_str(v: str, default_true: bool = True) -> str:
    if v is None or v == "":
        return "TRUE" if default_true else "FALSE"
    return "TRUE" if str(v).strip().upper() == "TRUE" else "FALSE"

def sanitize_name_for_folder(s: str) -> str:
    s = (s or "").strip()
    s = re.sub(r"[\\/:*?\"<>|]+", "_", s)
    s = re.sub(r"\s+", " ", s)
    return s[:120].strip() or "SinNombre"

# =========================
# GOOGLE AUTH
# =========================
@st.cache_resource
def get_creds():
    info = {
        "type": st.secrets["GCP_TYPE"],
        "project_id": st.secrets["GCP_PROJECT_ID"],
        "private_key_id": st.secrets["GCP_PRIVATE_KEY_ID"],
        "private_key": st.secrets["GCP_PRIVATE_KEY"],
        "client_email": st.secrets["GCP_CLIENT_EMAIL"],
        "client_id": st.secrets["GCP_CLIENT_ID"],
        "auth_uri": st.secrets.get("GCP_AUTH_URI", "https://accounts.google.com/o/oauth2/auth"),
        "token_uri": st.secrets.get("GCP_TOKEN_URI", "https://oauth2.googleapis.com/token"),
        "auth_provider_x509_cert_url": st.secrets.get(
            "GCP_AUTH_PROVIDER_X509_CERT_URL",
            "https://www.googleapis.com/oauth2/v1/certs",
        ),
        "client_x509_cert_url": st.secrets["GCP_CLIENT_X509_CERT_URL"],
    }
    scopes = [
        "https://www.googleapis.com/auth/spreadsheets",
        "https://www.googleapis.com/auth/drive",
    ]
    return Credentials.from_service_account_info(info, scopes=scopes)

@st.cache_resource
def authed_session() -> AuthorizedSession:
    return AuthorizedSession(get_creds())

# =========================
# API WRAPPERS (BACKOFF + CACHES)
# =========================
def _gs_url(path: str) -> str:
    return "https://sheets.googleapis.com" + path

def _drive_url(path: str) -> str:
    return "https://www.googleapis.com/drive/v3" + path

def _request_with_backoff(method: str, url: str, *, params=None, json_body=None, data=None, headers=None, timeout=30) -> dict:
    sess = authed_session()
    max_attempts = 7
    base_sleep = 0.8
    for attempt in range(1, max_attempts + 1):
        resp = sess.request(method, url, params=params, json=json_body, data=data, headers=headers, timeout=timeout)
        if resp.status_code == 429:
            sleep_s = (base_sleep * (2 ** (attempt - 1))) + (secrets.randbelow(250) / 1000.0)
            time.sleep(min(sleep_s, 10.0))
            continue
        if resp.status_code >= 400:
            raise RuntimeError(f"API error {resp.status_code}: {resp.text}")
        return resp.json() if resp.text else {}
    raise RuntimeError("API request failed after retries (429).")

@st.cache_data(ttl=12, show_spinner=False)
def sheets_get_cached(spreadsheet_id: str) -> dict:
    return _request_with_backoff("GET", _gs_url(f"/v4/spreadsheets/{spreadsheet_id}"))

@st.cache_data(ttl=12, show_spinner=False)
def sheets_values_get_cached(range_a1: str) -> List[List[str]]:
    enc = requests.utils.quote(range_a1, safe="")
    data = _request_with_backoff("GET", _gs_url(f"/v4/spreadsheets/{SHEET_ID}/values/{enc}"))
    return data.get("values", [])

def _invalidate_cache():
    st.cache_data.clear()

def sheets_values_update(range_a1: str, values: List[List[str]]):
    enc = requests.utils.quote(range_a1, safe="")
    _request_with_backoff(
        "PUT",
        _gs_url(f"/v4/spreadsheets/{SHEET_ID}/values/{enc}"),
        params={"valueInputOption": "RAW"},
        json_body={"range": range_a1, "majorDimension": "ROWS", "values": values},
    )
    _invalidate_cache()

def sheets_values_append(range_a1: str, values: List[List[str]]):
    enc = requests.utils.quote(range_a1, safe="")
    _request_with_backoff(
        "POST",
        _gs_url(f"/v4/spreadsheets/{SHEET_ID}/values/{enc}:append"),
        params={"valueInputOption": "RAW", "insertDataOption": "INSERT_ROWS"},
        json_body={"range": range_a1, "majorDimension": "ROWS", "values": values},
    )
    _invalidate_cache()

def sheets_batch_update(reqs: List[dict]):
    _request_with_backoff(
        "POST",
        _gs_url(f"/v4/spreadsheets/{SHEET_ID}:batchUpdate"),
        json_body={"requests": reqs},
    )
    _invalidate_cache()

def sheets_clear_range(range_a1: str):
    enc = requests.utils.quote(range_a1, safe="")
    _request_with_backoff("POST", _gs_url(f"/v4/spreadsheets/{SHEET_ID}/values/{enc}:clear"))
    _invalidate_cache()

# -------- Drive helpers --------
@st.cache_data(ttl=120, show_spinner=False)
def drive_find_folder_id(parent_id: str, name: str) -> Optional[str]:
    safe_name = (name or "").replace("'", "\\'")
    q = (
        "mimeType='application/vnd.google-apps.folder' and "
        f"name='{safe_name}' and "
        f"'{parent_id}' in parents and trashed=false"
    )
    data = _request_with_backoff("GET", _drive_url("/files"), params={"q": q, "fields": "files(id,name)"})
    files = data.get("files", [])
    return files[0]["id"] if files else None

def drive_create_folder(parent_id: str, name: str) -> str:
    body = {"name": name, "mimeType": "application/vnd.google-apps.folder", "parents": [parent_id]}
    data = _request_with_backoff("POST", _drive_url("/files"), params={"fields": "id"}, json_body=body)
    st.cache_data.clear()
    return data["id"]

def drive_get_or_create_folder(parent_id: str, name: str) -> str:
    name = sanitize_name_for_folder(name)
    fid = drive_find_folder_id(parent_id, name)
    if fid:
        return fid
    return drive_create_folder(parent_id, name)

def drive_upload_image_bytes(parent_folder_id: str, filename: str, content_type: str, file_bytes: bytes) -> Tuple[str, str]:
    sess = authed_session()
    metadata = {"name": filename, "parents": [parent_folder_id]}
    boundary = "===============" + secrets.token_hex(12)
    delimiter = f"\r\n--{boundary}\r\n"
    close_delim = f"\r\n--{boundary}--\r\n"

    multipart_body = (
        delimiter
        + "Content-Type: application/json; charset=UTF-8\r\n\r\n"
        + json.dumps(metadata)
        + delimiter
        + f"Content-Type: {content_type}\r\n\r\n"
    ).encode("utf-8") + file_bytes + close_delim.encode("utf-8")

    headers = {"Content-Type": f"multipart/related; boundary={boundary}"}
    url = "https://www.googleapis.com/upload/drive/v3/files?uploadType=multipart&fields=id,webViewLink"

    max_attempts = 7
    base_sleep = 0.8
    for attempt in range(1, max_attempts + 1):
        resp = sess.request("POST", url, data=multipart_body, headers=headers, timeout=60)
        if resp.status_code == 429:
            time.sleep(min((base_sleep * (2 ** (attempt - 1))) + (secrets.randbelow(250) / 1000.0), 10.0))
            continue
        if resp.status_code >= 400:
            raise RuntimeError(f"Drive upload error {resp.status_code}: {resp.text}")
        data = resp.json()
        st.cache_data.clear()
        return data["id"], data.get("webViewLink", "")
    raise RuntimeError("Drive upload failed after retries (429).")

def drive_download_bytes(file_id: str) -> bytes:
    # files.get alt=media
    sess = authed_session()
    url = f"https://www.googleapis.com/drive/v3/files/{file_id}"
    params = {"alt": "media"}
    max_attempts = 7
    base_sleep = 0.8
    for attempt in range(1, max_attempts + 1):
        resp = sess.request("GET", url, params=params, timeout=60)
        if resp.status_code == 429:
            time.sleep(min((base_sleep * (2 ** (attempt - 1))) + (secrets.randbelow(250) / 1000.0), 10.0))
            continue
        if resp.status_code >= 400:
            raise RuntimeError(f"Drive download error {resp.status_code}: {resp.text}")
        return resp.content
    raise RuntimeError("Drive download failed after retries (429).")

# =========================
# SHEETS: TABLE UTILITIES
# =========================
def ensure_tabs_and_headers(force_wipe: bool):
    meta = sheets_get_cached(SHEET_ID)
    existing = {s["properties"]["title"]: s["properties"]["sheetId"] for s in meta.get("sheets", [])}

    # Add missing tabs
    reqs = []
    for title in TABS_HEADERS.keys():
        if title not in existing:
            reqs.append({"addSheet": {"properties": {"title": title}}})
    if reqs:
        sheets_batch_update(reqs)
        meta = sheets_get_cached(SHEET_ID)
        existing = {s["properties"]["title"]: s["properties"]["sheetId"] for s in meta.get("sheets", [])}

    # Write headers (overwrite row 1)
    for tab, headers in TABS_HEADERS.items():
        sheets_values_update(f"{tab}!A1", [headers])
        if force_wipe:
            # clear from row 2 down (wipe data)
            sheets_clear_range(f"{tab}!A2:ZZ")

def read_table(tab: str) -> Tuple[List[str], List[Dict[str, str]]]:
    values = sheets_values_get_cached(f"{tab}!A1:ZZ")
    if not values:
        return [], []
    headers = [str(x).strip() for x in values[0]]
    rows: List[Dict[str, str]] = []
    for r in values[1:]:
        d: Dict[str, str] = {}
        for i, h in enumerate(headers):
            d[h] = str(r[i]).strip() if i < len(r) and r[i] is not None else ""
        if any(v != "" for v in d.values()):
            rows.append(d)
    return headers, rows

def write_table(tab: str, headers: List[str], rows: List[Dict[str, str]]):
    out: List[List[str]] = [headers]
    for d in rows:
        out.append([d.get(h, "") for h in headers])
    sheets_values_update(f"{tab}!A1", out)

# =========================
# AUTH: PASSWORD + INVITES
# =========================
def pbkdf2_hash_password(password: str, pepper: str, iterations: int = 210_000) -> str:
    salt = secrets.token_bytes(16)
    pw = (password + pepper).encode("utf-8")
    dk = hashlib.pbkdf2_hmac("sha256", pw, salt, iterations, dklen=32)
    return "pbkdf2_sha256$%d$%s$%s" % (
        iterations,
        base64.urlsafe_b64encode(salt).decode("utf-8"),
        base64.urlsafe_b64encode(dk).decode("utf-8"),
    )

def pbkdf2_verify_password(password: str, stored: str, pepper: str) -> bool:
    try:
        algo, it_s, salt_b64, hash_b64 = stored.split("$", 3)
        if algo != "pbkdf2_sha256":
            return False
        iterations = int(it_s)
        salt = base64.urlsafe_b64decode(salt_b64.encode("utf-8"))
        expected = base64.urlsafe_b64decode(hash_b64.encode("utf-8"))
        pw = (password + pepper).encode("utf-8")
        dk = hashlib.pbkdf2_hmac("sha256", pw, salt, iterations, dklen=32)
        return hmac.compare_digest(dk, expected)
    except Exception:
        return False

def hash_invite_code(code: str, pepper: str) -> str:
    h = hashlib.sha256((code + pepper).encode("utf-8")).hexdigest()
    return f"sha256${h}"

def verify_invite_code(code: str, stored: str, pepper: str) -> bool:
    try:
        algo, hexhash = stored.split("$", 1)
        if algo != "sha256":
            return False
        h = hashlib.sha256((code + pepper).encode("utf-8")).hexdigest()
        return hmac.compare_digest(h, hexhash)
    except Exception:
        return False

def bootstrap_admin_users():
    _, users = read_table("Users")
    changed = False
    existing_map = {norm_email(u.get("email", "")): u for u in users if u.get("email")}

    for em in BOOTSTRAP_ADMIN_EMAILS:
        if em not in existing_map:
            users.append({
                "user_id": f"USR-{secrets.token_hex(4)}",
                "email": em,
                "full_name": "Admin",
                "is_admin": "TRUE",
                "is_active": "TRUE",
                "password_hash": "",
                "created_at": iso(now_utc()),
                "last_login_at": "",
            })
            changed = True
        else:
            u = existing_map[em]
            if u.get("is_admin", "").upper() != "TRUE":
                u["is_admin"] = "TRUE"; changed = True
            if u.get("is_active", "").upper() != "TRUE":
                u["is_active"] = "TRUE"; changed = True

    if changed:
        write_table("Users", TABS_HEADERS["Users"], users)

def get_user_by_email(email: str) -> Optional[Dict[str, str]]:
    _, users = read_table("Users")
    em = norm_email(email)
    for u in users:
        if norm_email(u.get("email", "")) == em:
            return u
    return None

def upsert_user(user: Dict[str, str]):
    _, users = read_table("Users")
    em = norm_email(user.get("email", ""))
    found = False
    for i, u in enumerate(users):
        if norm_email(u.get("email", "")) == em:
            users[i] = {**u, **user}
            found = True
            break
    if not found:
        users.append(user)
    write_table("Users", TABS_HEADERS["Users"], users)

def generate_invite(email: str, created_by: str) -> str:
    code = "".join(str(secrets.randbelow(10)) for _ in range(6))
    expires = now_utc() + timedelta(hours=INVITE_EXPIRY_HOURS)
    row = {
        "email": norm_email(email),
        "invite_code_hash": hash_invite_code(code, APP_PEPPER),
        "expires_at": iso(expires),
        "used_at": "",
        "created_by": norm_email(created_by),
        "created_at": iso(now_utc()),
    }
    sheets_values_append("Invites!A1", [[row.get(h, "") for h in TABS_HEADERS["Invites"]]])
    return code

def find_valid_invite(email: str, code: str) -> Optional[Dict[str, str]]:
    _, invites = read_table("Invites")
    em = norm_email(email)
    for inv in reversed(invites):
        if norm_email(inv.get("email", "")) != em:
            continue
        if inv.get("used_at"):
            continue
        exp = parse_iso(inv.get("expires_at", ""))
        if not exp or now_utc() > exp:
            continue
        if verify_invite_code(code, inv.get("invite_code_hash", ""), APP_PEPPER):
            return inv
    return None

def mark_invite_used(invite: Dict[str, str]):
    _, invites = read_table("Invites")
    em = norm_email(invite.get("email", ""))
    ca = invite.get("created_at", "")
    for i, inv in enumerate(invites):
        if norm_email(inv.get("email", "")) == em and inv.get("created_at", "") == ca:
            invites[i]["used_at"] = iso(now_utc())
            break
    write_table("Invites", TABS_HEADERS["Invites"], invites)

@dataclass
class AuthUser:
    email: str
    full_name: str
    is_admin: bool

def set_auth(user: AuthUser):
    st.session_state["auth"] = {"email": user.email, "full_name": user.full_name, "is_admin": user.is_admin}

def get_auth() -> Optional[AuthUser]:
    d = st.session_state.get("auth")
    if not d:
        return None
    return AuthUser(email=d["email"], full_name=d.get("full_name", ""), is_admin=bool(d.get("is_admin")))

def clear_auth():
    st.session_state.pop("auth", None)
    st.session_state.pop("selected_community_id", None)
    st.session_state.pop("selected_report_id", None)

# =========================
# ACCESS / PERMISSIONS
# =========================
def list_user_access(email: str) -> List[Dict[str, str]]:
    _, rows = read_table("UserCommunityAccess")
    em = norm_email(email)
    out = []
    for r in rows:
        if norm_email(r.get("email", "")) != em:
            continue
        if safe_bool_str(r.get("is_active", ""), True) != "TRUE":
            continue
        out.append({
            "email": em,
            "community_id": r.get("community_id", ""),
            "role": (r.get("role", "") or "viewer").lower(),
            "can_create_reports": safe_bool_str(r.get("can_create_reports", ""), False) == "TRUE",
            "can_view_summary": safe_bool_str(r.get("can_view_summary", ""), False) == "TRUE",
        })
    return out

def grant_access(email: str, community_id: str, role: str, can_create_reports: bool, can_view_summary: bool):
    headers, rows = read_table("UserCommunityAccess")
    em = norm_email(email)
    role = (role or "viewer").strip().lower()
    if role not in ROLES:
        role = "viewer"

    found = False
    for i, r in enumerate(rows):
        if norm_email(r.get("email", "")) == em and r.get("community_id", "") == community_id:
            rows[i]["role"] = role
            rows[i]["is_active"] = "TRUE"
            rows[i]["can_create_reports"] = "TRUE" if can_create_reports else "FALSE"
            rows[i]["can_view_summary"] = "TRUE" if can_view_summary else "FALSE"
            found = True
            break

    if not found:
        rows.append({
            "email": em,
            "community_id": community_id,
            "role": role,
            "is_active": "TRUE",
            "can_create_reports": "TRUE" if can_create_reports else "FALSE",
            "can_view_summary": "TRUE" if can_view_summary else "FALSE",
        })

    write_table("UserCommunityAccess", TABS_HEADERS["UserCommunityAccess"], rows)

def revoke_access(email: str, community_id: str):
    headers, rows = read_table("UserCommunityAccess")
    em = norm_email(email)
    for i, r in enumerate(rows):
        if norm_email(r.get("email", "")) == em and r.get("community_id", "") == community_id:
            rows[i]["is_active"] = "FALSE"
            break
    write_table("UserCommunityAccess", TABS_HEADERS["UserCommunityAccess"], rows)

# =========================
# COMMUNITIES / INSTALLATIONS / TASKS
# =========================
def list_communities(include_inactive: bool = False) -> List[Dict[str, str]]:
    _, comms = read_table("Communities")
    out = []
    for c in comms:
        active = safe_bool_str(c.get("is_active", ""), True)
        if include_inactive or active == "TRUE":
            if c.get("community_id"):
                out.append({**c, "is_active": active})
    out.sort(key=lambda x: x.get("community_name", ""))
    return out

def create_community(name: str) -> Dict[str, str]:
    name = name.strip()
    _, comms = read_table("Communities")
    existing_ids = [c.get("community_id", "") for c in comms if c.get("community_id", "").startswith("COM-")]
    nums = []
    for cid in existing_ids:
        try:
            nums.append(int(cid.split("-")[1]))
        except Exception:
            pass
    next_n = (max(nums) + 1) if nums else 1
    cid = f"COM-{next_n:04d}"
    row = {
        "community_id": cid,
        "community_name": name,
        "is_active": "TRUE",
        "cover_file_id": "",
        "cover_web_view_link": "",
        "created_at": iso(now_utc()),
    }
    sheets_values_append("Communities!A1", [[row.get(h, "") for h in TABS_HEADERS["Communities"]]])
    return row

def set_community_active(community_id: str, active: bool):
    _, comms = read_table("Communities")
    for i, c in enumerate(comms):
        if c.get("community_id") == community_id:
            comms[i]["is_active"] = "TRUE" if active else "FALSE"
            break
    write_table("Communities", TABS_HEADERS["Communities"], comms)

def update_community_cover(community_id: str, file_id: str, web_link: str):
    _, comms = read_table("Communities")
    for i, c in enumerate(comms):
        if c.get("community_id") == community_id:
            comms[i]["cover_file_id"] = file_id
            comms[i]["cover_web_view_link"] = web_link
            break
    write_table("Communities", TABS_HEADERS["Communities"], comms)

def list_installations(community_id: str, include_inactive: bool = False) -> List[Dict[str, str]]:
    _, rows = read_table("Installations")
    out = []
    for r in rows:
        if r.get("community_id") != community_id:
            continue
        active = safe_bool_str(r.get("is_active", ""), True)
        if include_inactive or active == "TRUE":
            out.append({**r, "is_active": active})
    out.sort(key=lambda x: (x.get("category", ""), x.get("installation_name", "")))
    return out

def add_installation(community_id: str, category: str, name: str) -> Dict[str, str]:
    row = {
        "installation_id": f"INS-{secrets.token_hex(4)}",
        "community_id": community_id,
        "category": category.strip(),
        "installation_name": name.strip(),
        "is_active": "TRUE",
        "created_at": iso(now_utc()),
    }
    sheets_values_append("Installations!A1", [[row.get(h, "") for h in TABS_HEADERS["Installations"]]])
    return row

def set_installation_active(installation_id: str, active: bool):
    _, rows = read_table("Installations")
    for i, r in enumerate(rows):
        if r.get("installation_id") == installation_id:
            rows[i]["is_active"] = "TRUE" if active else "FALSE"
            break
    write_table("Installations", TABS_HEADERS["Installations"], rows)

def list_tasks(community_id: str, installation_id: Optional[str] = None, include_inactive: bool = False) -> List[Dict[str, str]]:
    _, rows = read_table("Tasks")
    out = []
    for r in rows:
        if r.get("community_id") != community_id:
            continue
        if installation_id and r.get("installation_id") != installation_id:
            continue
        active = safe_bool_str(r.get("is_active", ""), True)
        if include_inactive or active == "TRUE":
            # normalize order int-ish
            o = r.get("task_order", "")
            try:
                ordv = int(float(o))
            except Exception:
                ordv = 9999
            out.append({**r, "is_active": active, "_order": ordv})
    out.sort(key=lambda x: (x.get("installation_id", ""), x.get("_order", 9999), x.get("task_name", "")))
    return out

def add_task(community_id: str, installation_id: str, task_name: str, task_order: int, is_required: bool) -> Dict[str, str]:
    row = {
        "task_id": f"TSK-{secrets.token_hex(4)}",
        "community_id": community_id,
        "installation_id": installation_id,
        "task_name": task_name.strip(),
        "task_order": str(int(task_order)),
        "is_required": "TRUE" if is_required else "FALSE",
        "is_active": "TRUE",
        "created_at": iso(now_utc()),
    }
    sheets_values_append("Tasks!A1", [[row.get(h, "") for h in TABS_HEADERS["Tasks"]]])
    return row

def set_task_active(task_id: str, active: bool):
    _, rows = read_table("Tasks")
    for i, r in enumerate(rows):
        if r.get("task_id") == task_id:
            rows[i]["is_active"] = "TRUE" if active else "FALSE"
            break
    write_table("Tasks", TABS_HEADERS["Tasks"], rows)

def seed_default_installations_and_tasks_if_empty(community_id: str):
    inst_all = list_installations(community_id, include_inactive=True)
    tasks_all = list_tasks(community_id, include_inactive=True)
    if inst_all or tasks_all:
        return

    inst_map = {}  # (category,name)->installation_id
    for cat, inst_name, tasks in DEFAULT_INSTALLATIONS_AND_TASKS:
        inst = add_installation(community_id, cat, inst_name)
        inst_map[(cat, inst_name)] = inst["installation_id"]
        for tname, order, req in tasks:
            add_task(community_id, inst["installation_id"], tname, order, req)

# =========================
# DRIVE FOLDERS
# =========================
def ensure_drive_ready():
    if not FOTOS_ROOT_FOLDER_ID:
        st.error("Falta FOTOS_ROOT_FOLDER_ID en secrets.toml")
        st.stop()

def get_community_folder_id(community_id: str, community_name: str) -> str:
    ensure_drive_ready()
    folder_name = f"{community_id}__{sanitize_name_for_folder(community_name)}"
    return drive_get_or_create_folder(FOTOS_ROOT_FOLDER_ID, folder_name)

def get_report_folder_id(comm_folder_id: str, report_id: str) -> str:
    return drive_get_or_create_folder(comm_folder_id, f"Reports__{report_id}")

def get_installation_folder_id(report_folder_id: str, installation_id: str, installation_name: str) -> str:
    folder_name = f"{installation_id}__{sanitize_name_for_folder(installation_name)}"
    return drive_get_or_create_folder(report_folder_id, folder_name)

# =========================
# REPORTS + ITEMS
# =========================
def list_reports_for_community(community_id: str) -> List[Dict[str, str]]:
    _, rows = read_table("Reports")
    out = [r for r in rows if r.get("community_id") == community_id]
    out.sort(key=lambda x: (x.get("report_date", ""), x.get("created_at", "")), reverse=True)
    return out

def list_reports_for_user_and_community(email: str, community_id: str) -> List[Dict[str, str]]:
    _, rows = read_table("Reports")
    em = norm_email(email)
    out = [r for r in rows if r.get("community_id") == community_id and norm_email(r.get("created_by", "")) == em]
    out.sort(key=lambda x: x.get("updated_at", "") or x.get("created_at", ""), reverse=True)
    return out

def get_report(report_id: str) -> Optional[Dict[str, str]]:
    _, rows = read_table("Reports")
    for r in rows:
        if r.get("report_id") == report_id:
            return r
    return None

def upsert_report(report: Dict[str, str]):
    _, rows = read_table("Reports")
    rid = report.get("report_id", "")
    found = False
    for i, r in enumerate(rows):
        if r.get("report_id") == rid:
            rows[i] = {**r, **report}
            found = True
            break
    if not found:
        rows.append(report)
    write_table("Reports", TABS_HEADERS["Reports"], rows)

def create_new_report(community_id: str, community_name: str, created_by: str, report_date: date) -> Dict[str, str]:
    rid = f"RPT-{secrets.token_hex(5)}"
    row = {
        "report_id": rid,
        "community_id": community_id,
        "community_name": community_name,
        "report_date": report_date.isoformat(),
        "status": "Draft",
        "created_by": norm_email(created_by),
        "created_at": iso(now_utc()),
        "updated_at": iso(now_utc()),
    }
    sheets_values_append("Reports!A1", [[row.get(h, "") for h in TABS_HEADERS["Reports"]]])
    return row

def list_report_items(report_id: str) -> List[Dict[str, str]]:
    _, rows = read_table("ReportItems")
    out = [r for r in rows if r.get("report_id") == report_id]
    # order by category/installation/task_order
    def ordv(x):
        try:
            return int(float(x.get("task_order", "9999") or 9999))
        except Exception:
            return 9999
    out.sort(key=lambda x: (x.get("category", ""), x.get("installation_name", ""), ordv(x), x.get("task_name", "")))
    return out

def get_report_items_map(report_id: str) -> Dict[Tuple[str, str], Dict[str, str]]:
    # (installation_id, task_id) -> row
    items = list_report_items(report_id)
    m = {}
    for r in items:
        m[(r.get("installation_id", ""), r.get("task_id", ""))] = r
    return m

def ensure_report_items_from_catalog(report: Dict[str, str], actor_email: str):
    """
    Crea items faltantes en ReportItems según Tasks activas de la comunidad.
    No borra items antiguos (por si cambió catálogo), solo completa los que falten.
    """
    community_id = report["community_id"]
    inst = list_installations(community_id, include_inactive=False)
    inst_map = {i["installation_id"]: i for i in inst}

    tasks = list_tasks(community_id, include_inactive=False)
    tasks_by_inst: Dict[str, List[Dict[str, str]]] = {}
    for t in tasks:
        tasks_by_inst.setdefault(t["installation_id"], []).append(t)

    existing = get_report_items_map(report["report_id"])
    new_rows = []

    for ins in inst:
        ins_id = ins["installation_id"]
        for t in tasks_by_inst.get(ins_id, []):
            key = (ins_id, t["task_id"])
            if key in existing:
                continue
            new_rows.append({
                "report_item_id": f"RPI-{secrets.token_hex(5)}",
                "report_id": report["report_id"],
                "community_id": community_id,
                "installation_id": ins_id,
                "installation_name": ins["installation_name"],
                "category": ins["category"],
                "task_id": t["task_id"],
                "task_name": t["task_name"],
                "task_order": t.get("task_order", "9999"),
                "status": "pending",
                "note": "",
                "photo_file_id": "",
                "photo_web_view_link": "",
                "updated_at": iso(now_utc()),
                "updated_by": norm_email(actor_email),
            })

    if new_rows:
        for row in new_rows:
            sheets_values_append("ReportItems!A1", [[row.get(h, "") for h in TABS_HEADERS["ReportItems"]]])

def upsert_report_item(item: Dict[str, str]):
    _, rows = read_table("ReportItems")
    rid = item.get("report_id", "")
    ins_id = item.get("installation_id", "")
    task_id = item.get("task_id", "")
    found = False
    for i, r in enumerate(rows):
        if r.get("report_id") == rid and r.get("installation_id") == ins_id and r.get("task_id") == task_id:
            rows[i] = {**r, **item}
            found = True
            break
    if not found:
        if not item.get("report_item_id"):
            item["report_item_id"] = f"RPI-{secrets.token_hex(5)}"
        rows.append(item)
    write_table("ReportItems", TABS_HEADERS["ReportItems"], rows)

def set_report_status(report_id: str, status: str):
    rep = get_report(report_id)
    if not rep:
        return
    rep["status"] = status
    rep["updated_at"] = iso(now_utc())
    upsert_report(rep)

# =========================
# EXPORTERS (PDF / WORD)
# =========================
def build_report_structure(report: Dict[str, str]) -> Dict[str, Dict[str, List[Dict[str, str]]]]:
    """
    returns: {category: {installation_name: [items...]}}
    """
    items = list_report_items(report["report_id"])
    tree: Dict[str, Dict[str, List[Dict[str, str]]]] = {}
    for it in items:
        cat = it.get("category", "Sin categoría") or "Sin categoría"
        inst_name = it.get("installation_name", "Sin instalación") or "Sin instalación"
        tree.setdefault(cat, {}).setdefault(inst_name, []).append(it)

    # order tasks within installation by task_order
    def ord_task(x):
        try:
            return int(float(x.get("task_order", "9999") or 9999))
        except Exception:
            return 9999

    for cat in tree:
        for ins in tree[cat]:
            tree[cat][ins].sort(key=lambda x: (ord_task(x), x.get("task_name", "")))
    return tree

def export_pdf(report: Dict[str, str]) -> bytes:
    if not REPORTLAB_OK:
        raise RuntimeError("reportlab no está instalado. Agrega reportlab en requirements.")
    tree = build_report_structure(report)

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    margin = 36
    y = height - margin

    def draw_title(txt, size=14):
        nonlocal y
        c.setFont("Helvetica-Bold", size)
        c.drawString(margin, y, txt)
        y -= 18

    def draw_text(txt, size=10):
        nonlocal y
        c.setFont("Helvetica", size)
        # simple wrap
        max_chars = 110
        for line in (txt or "").split("\n"):
            while len(line) > max_chars:
                c.drawString(margin, y, line[:max_chars])
                y -= 14
                line = line[max_chars:]
            c.drawString(margin, y, line)
            y -= 14

    draw_title(f"INFORME — {report.get('community_name','')} ({report.get('community_id','')})", 16)
    draw_text(f"Fecha: {report.get('report_date','')}   Estado: {report.get('status','')}")
    draw_text(f"Creado por: {report.get('created_by','')}   Última actualización: {report.get('updated_at','')}")
    y -= 6

    # 3-column layout widths
    col1_w = 170
    col2_w = 230
    col3_w = width - margin*2 - col1_w - col2_w
    row_h = 58

    def new_page_if_needed(min_space=80):
        nonlocal y
        if y < margin + min_space:
            c.showPage()
            y = height - margin

    for cat in sorted(tree.keys(), key=lambda s: (DEFAULT_CATEGORIES.index(s) if s in DEFAULT_CATEGORIES else 999, s)):
        new_page_if_needed()
        draw_title(f"Área: {cat}", 13)
        for inst_name, items in tree[cat].items():
            new_page_if_needed()
            c.setFont("Helvetica-Bold", 11)
            c.drawString(margin, y, f"Instalación: {inst_name}")
            y -= 16

            # table header
            c.setFont("Helvetica-Bold", 9)
            x1 = margin
            x2 = margin + col1_w
            x3 = margin + col1_w + col2_w
            c.drawString(x1, y, "Tarea")
            c.drawString(x2, y, "Estado / Observación")
            c.drawString(x3, y, "Foto")
            y -= 10

            for it in items:
                new_page_if_needed(min_space=row_h + 30)
                # background red soft for fail
                status = (it.get("status") or "pending").lower()
                if status == "fail":
                    c.setFillColorRGB(1.0, 0.90, 0.90)
                    c.rect(margin, y - row_h + 12, width - margin*2, row_h, stroke=0, fill=1)
                    c.setFillColorRGB(0, 0, 0)

                # borders
                c.rect(margin, y - row_h + 12, width - margin*2, row_h, stroke=1, fill=0)
                c.line(x2, y - row_h + 12, x2, y + 12)
                c.line(x3, y - row_h + 12, x3, y + 12)

                # col1: task
                c.setFont("Helvetica-Bold", 9)
                c.drawString(x1 + 6, y, it.get("task_name", "")[:40])

                # col2: status + note
                c.setFont("Helvetica", 9)
                badge = "OK" if status == "ok" else ("FALLA" if status == "fail" else "PEND.")
                c.drawString(x2 + 6, y, f"[{badge}] {it.get('note','')[:55]}")

                # col3: photo (thumbnail if exists)
                file_id = it.get("photo_file_id", "")
                if file_id:
                    try:
                        img_bytes = drive_download_bytes(file_id)
                        img = ImageReader(io.BytesIO(img_bytes))
                        thumb_w = col3_w - 16
                        thumb_h = row_h - 18
                        c.drawImage(img, x3 + 8, y - row_h + 20, width=thumb_w, height=thumb_h, preserveAspectRatio=True, anchor='c')
                    except Exception:
                        c.setFont("Helvetica", 8)
                        c.drawString(x3 + 6, y, "Foto (error al cargar)")

                y -= row_h + 10

            y -= 8

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.getvalue()

def export_docx(report: Dict[str, str]) -> bytes:
    if not PYDOCX_OK:
        raise RuntimeError("python-docx no está instalado. Agrega python-docx en requirements.")
    tree = build_report_structure(report)

    doc = Document()
    doc.add_heading(f"INFORME — {report.get('community_name','')} ({report.get('community_id','')})", level=1)
    doc.add_paragraph(f"Fecha: {report.get('report_date','')}   Estado: {report.get('status','')}")
    doc.add_paragraph(f"Creado por: {report.get('created_by','')}   Última actualización: {report.get('updated_at','')}")

    for cat in sorted(tree.keys(), key=lambda s: (DEFAULT_CATEGORIES.index(s) if s in DEFAULT_CATEGORIES else 999, s)):
        doc.add_heading(f"Área: {cat}", level=2)
        for inst_name, items in tree[cat].items():
            doc.add_heading(f"Instalación: {inst_name}", level=3)

            table = doc.add_table(rows=1, cols=3)
            hdr = table.rows[0].cells
            hdr[0].text = "Tarea"
            hdr[1].text = "Estado / Observación"
            hdr[2].text = "Foto"

            for it in items:
                row = table.add_row().cells
                status = (it.get("status") or "pending").lower()
                badge = "OK" if status == "ok" else ("FALLA" if status == "fail" else "PEND.")
                row[0].text = it.get("task_name", "")
                row[1].text = f"[{badge}] {it.get('note','')}"
                if it.get("photo_file_id"):
                    try:
                        img_bytes = drive_download_bytes(it["photo_file_id"])
                        # docx needs a file-like object
                        row[2].paragraphs[0].add_run().add_picture(io.BytesIO(img_bytes), width=Inches(2.0))
                    except Exception:
                        row[2].text = "(error al cargar foto)"
                else:
                    row[2].text = ""

            doc.add_paragraph("")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()

# =========================
# BOOT / RESET (TESTING)
# =========================
def boot(force_wipe: bool):
    ensure_tabs_and_headers(force_wipe=force_wipe)
    bootstrap_admin_users()

# =========================
# UI COMPONENTS
# =========================
def status_badge(status: str) -> str:
    if status == "ok":
        return "🟢 OK"
    if status == "fail":
        return "🔴 FALLA"
    return "⚪ PEND."

def status_bg(status: str) -> str:
    if status == "fail":
        return "#fee2e2"  # soft red
    return "#ffffff"

# =========================
# BOOT EXEC
# =========================
if "boot_ok" not in st.session_state:
    # En testing: NO borramos por defecto, pero tendrás un botón Admin para reset.
    try:
        boot(force_wipe=False)
        st.session_state["boot_ok"] = True
    except Exception as e:
        st.error("No pude inicializar la estructura del Google Sheet.")
        st.exception(e)
        st.stop()

# =========================
# LOGIN
# =========================
auth = get_auth()
st.title("🛡️ Control Comunidades")

if not auth:
    st.info("Ingreso con contraseña o primer acceso con código. (Admins bootstrap requieren activación inicial)")

    with st.expander("🧰 Activación inicial (solo admins bootstrap)"):
        b_email = st.text_input("Email admin bootstrap", placeholder="bnbpartnerscommunity@gmail.com")
        b_token = st.text_input("Token secreto", type="password", placeholder="BOOTSTRAP_SETUP_TOKEN")
        b_pass1 = st.text_input("Nueva contraseña", type="password", key="bpass1")
        b_pass2 = st.text_input("Repetir contraseña", type="password", key="bpass2")

        if st.button("Activar admin"):
            em = norm_email(b_email)
            if em not in BOOTSTRAP_ADMIN_EMAILS:
                st.error("Ese correo no está autorizado como admin bootstrap.")
            elif not BOOTSTRAP_SETUP_TOKEN:
                st.error("Falta BOOTSTRAP_SETUP_TOKEN en Secrets.")
            elif b_token != BOOTSTRAP_SETUP_TOKEN:
                st.error("Token secreto incorrecto.")
            elif b_pass1 != b_pass2:
                st.error("Las contraseñas no coinciden.")
            elif len(b_pass1) < 8 or not any(c.isdigit() for c in b_pass1) or not any(c.isalpha() for c in b_pass1):
                st.error("Contraseña débil: mínimo 8, al menos 1 letra y 1 número.")
            else:
                u = get_user_by_email(em)
                if not u:
                    st.error("Usuario bootstrap no existe en Users.")
                else:
                    u["password_hash"] = pbkdf2_hash_password(b_pass1, APP_PEPPER)
                    u["last_login_at"] = iso(now_utc())
                    upsert_user(u)
                    st.success("Admin activado. Ahora puedes ingresar con contraseña.")

    c1, c2 = st.columns(2, gap="large")

    with c1:
        st.subheader("Ingreso con contraseña")
        email = st.text_input("Email", key="login_email", placeholder="usuario@empresa.com")
        password = st.text_input("Contraseña", type="password", key="login_password")
        if st.button("Ingresar", use_container_width=True):
            em = norm_email(email)
            u = get_user_by_email(em) if EMAIL_RE.match(em) else None
            if not u or (safe_bool_str(u.get("is_active", ""), True) != "TRUE"):
                st.error("Usuario no existe o está inactivo.")
            else:
                stored = u.get("password_hash", "")
                if not stored:
                    st.error("Este usuario no tiene contraseña aún. Debe ingresar con código.")
                elif pbkdf2_verify_password(password, stored, APP_PEPPER):
                    u["last_login_at"] = iso(now_utc())
                    upsert_user(u)
                    set_auth(AuthUser(
                        email=em,
                        full_name=u.get("full_name", "") or em,
                        is_admin=(safe_bool_str(u.get("is_admin", ""), False) == "TRUE")
                    ))
                    st.rerun()
                else:
                    st.error("Contraseña incorrecta.")

    with c2:
        st.subheader("Primer acceso con código")
        email2 = st.text_input("Email", key="invite_email", placeholder="usuario@empresa.com")
        code = st.text_input("Código (6 dígitos)", key="invite_code", max_chars=6)
        if st.button("Validar código", use_container_width=True):
            em = norm_email(email2)
            u = get_user_by_email(em) if EMAIL_RE.match(em) else None
            if not u or (safe_bool_str(u.get("is_active", ""), True) != "TRUE"):
                st.error("Usuario no existe o está inactivo.")
            else:
                inv = find_valid_invite(em, code.strip())
                if not inv:
                    st.error("Código incorrecto, vencido o ya usado.")
                else:
                    st.session_state["pending_set_password_email"] = em
                    st.session_state["pending_invite_created_at"] = inv.get("created_at", "")
                    st.success("Código validado. Ahora define tu contraseña abajo.")
                    st.rerun()

    pending_email = st.session_state.get("pending_set_password_email")
    if pending_email:
        st.divider()
        st.subheader("✅ Definir contraseña (primer acceso)")
        p1 = st.text_input("Nueva contraseña", type="password", key="new_pass_1")
        p2 = st.text_input("Repetir contraseña", type="password", key="new_pass_2")

        def strong_enough(p: str) -> bool:
            return len(p) >= 8 and any(ch.isdigit() for ch in p) and any(ch.isalpha() for ch in p)

        if st.button("Guardar contraseña", use_container_width=True):
            if p1 != p2:
                st.error("Las contraseñas no coinciden.")
            elif not strong_enough(p1):
                st.error("Contraseña débil. Mínimo 8, al menos 1 letra y 1 número.")
            else:
                _, invites = read_table("Invites")
                created_at = st.session_state.get("pending_invite_created_at", "")
                inv_row = None
                for inv in reversed(invites):
                    if norm_email(inv.get("email", "")) == pending_email and inv.get("created_at", "") == created_at:
                        inv_row = inv
                        break
                if inv_row:
                    mark_invite_used(inv_row)

                u = get_user_by_email(pending_email)
                u["password_hash"] = pbkdf2_hash_password(p1, APP_PEPPER)
                u["last_login_at"] = iso(now_utc())
                upsert_user(u)

                st.session_state.pop("pending_set_password_email", None)
                st.session_state.pop("pending_invite_created_at", None)

                set_auth(AuthUser(
                    email=pending_email,
                    full_name=u.get("full_name", "") or pending_email,
                    is_admin=(safe_bool_str(u.get("is_admin", ""), False) == "TRUE")
                ))
                st.success("Contraseña guardada. Sesión iniciada.")
                st.rerun()

    st.stop()

# =========================
# AUTHENTICATED UI
# =========================
st.sidebar.success(f"Conectado: {auth.full_name}\n\n{auth.email}")
if st.sidebar.button("Cerrar sesión"):
    clear_auth()
    st.rerun()

# Reset DB (testing) - admin only
if auth.is_admin:
    with st.sidebar.expander("🧨 Testing: Reset DB (pisa datos)"):
        st.caption("Esto reescribe headers y borra todas las filas (desde la 2) en todas las tabs del esquema nuevo.")
        confirm = st.text_input("Escribe CONFIRMAR para ejecutar", key="reset_confirm")
        if st.button("RESET AHORA", type="primary"):
            if confirm.strip().upper() != "CONFIRMAR":
                st.error("Debes escribir CONFIRMAR.")
            else:
                try:
                    boot(force_wipe=True)
                    st.success("Reset completado ✅")
                    st.rerun()
                except Exception as e:
                    st.error("No pude resetear.")
                    st.exception(e)

# Determine modules allowed
access_rows = list_user_access(auth.email)
comms = list_communities(include_inactive=False)
comm_map = {c["community_id"]: c for c in comms}

allowed_comms = []
can_create_reports_any = False
can_view_summary_any = False

for a in access_rows:
    cid = a["community_id"]
    if cid in comm_map:
        allowed_comms.append({
            "community_id": cid,
            "community_name": comm_map[cid]["community_name"],
            "cover_link": comm_map[cid].get("cover_web_view_link", ""),
            "role": a["role"],
            "can_create_reports": a["can_create_reports"],
            "can_view_summary": a["can_view_summary"],
        })
        can_create_reports_any = can_create_reports_any or a["can_create_reports"]
        can_view_summary_any = can_view_summary_any or a["can_view_summary"]

# Admin sees all communities
if auth.is_admin:
    for cid, c in comm_map.items():
        if not any(x["community_id"] == cid for x in allowed_comms):
            allowed_comms.append({
                "community_id": cid,
                "community_name": c["community_name"],
                "cover_link": c.get("cover_web_view_link", ""),
                "role": "admin",
                "can_create_reports": True,
                "can_view_summary": True,
            })
    can_create_reports_any = True
    can_view_summary_any = True

allowed_comms.sort(key=lambda x: x["community_name"])

modules = []
if auth.is_admin:
    modules.append("👤 Usuarios")
    modules.append("🏢 Comunidades")
if can_create_reports_any:
    modules.append("🧾 Informes")
if can_view_summary_any:
    modules.append("📊 Resumen")

if not modules:
    st.warning("No tienes módulos habilitados. Pide al admin que te asigne permisos.")
    st.stop()

module = st.sidebar.radio("Módulos", modules, index=0)

# =========================
# MODULE: USERS (ADMIN)
# =========================
if module == "👤 Usuarios":
    st.header("👤 Módulo Usuarios (Admin)")

    _, users = read_table("Users")
    users = sorted(users, key=lambda u: u.get("email", ""))
    st.dataframe(
        [{
            "Email": u.get("email", ""),
            "Nombre": u.get("full_name", ""),
            "Admin": safe_bool_str(u.get("is_admin", ""), False) == "TRUE",
            "Activo": safe_bool_str(u.get("is_active", ""), True) == "TRUE",
            "Último login": u.get("last_login_at", ""),
        } for u in users],
        use_container_width=True,
        hide_index=True
    )

    st.divider()
    st.subheader("Crear/actualizar usuario + asignar permisos + generar código")

    email = st.text_input("Email", placeholder="persona@empresa.com")
    full_name = st.text_input("Nombre completo", placeholder="Juan Pérez")
    is_admin_flag = st.checkbox("¿Es admin global?", value=False)

    comm_labels = [f"{c['community_name']} ({c['community_id']})" for c in comms]
    selected = st.multiselect("Comunidades", options=comm_labels)

    role = st.selectbox("Rol (por comunidad)", options=ROLES, index=1)
    can_create_reports = st.checkbox("Permiso: crear informes (Módulo Informes)", value=True)
    can_view_summary = st.checkbox("Permiso: ver resumen (Módulo Resumen)", value=False)

    if st.button("Guardar usuario y generar código"):
        em = norm_email(email)
        if not EMAIL_RE.match(em):
            st.error("Email inválido.")
        elif not full_name.strip():
            st.error("Falta nombre.")
        else:
            u = get_user_by_email(em)
            if not u:
                u = {
                    "user_id": f"USR-{secrets.token_hex(4)}",
                    "email": em,
                    "full_name": full_name.strip(),
                    "is_admin": "TRUE" if is_admin_flag else "FALSE",
                    "is_active": "TRUE",
                    "password_hash": "",
                    "created_at": iso(now_utc()),
                    "last_login_at": "",
                }
            else:
                # bootstrap users cannot be deactivated here (we don't show a deactivate button in this quick form)
                u["full_name"] = full_name.strip()
                u["is_admin"] = "TRUE" if is_admin_flag else "FALSE"
                u["is_active"] = "TRUE"

            upsert_user(u)

            # Assign access
            label_to_id = {f"{c['community_name']} ({c['community_id']})": c["community_id"] for c in comms}
            for lbl in selected:
                cid = label_to_id[lbl]
                grant_access(em, cid, role, can_create_reports, can_view_summary)

            code = generate_invite(em, auth.email)
            st.success("Listo. Envía este código al usuario para su primer ingreso:")
            st.code(code)

    st.divider()
    st.subheader("Activar/Desactivar usuario (no bootstrap)")
    non_bootstrap = [u for u in users if norm_email(u.get("email", "")) not in BOOTSTRAP_ADMIN_EMAILS]
    if non_bootstrap:
        choices = [u["email"] for u in non_bootstrap]
        target = st.selectbox("Usuario", options=choices)
        tu = get_user_by_email(target)
        active = safe_bool_str(tu.get("is_active", ""), True) == "TRUE"
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Desactivar" if active else "Activar"):
                tu["is_active"] = "FALSE" if active else "TRUE"
                upsert_user(tu)
                st.success("Actualizado.")
                st.rerun()
        with col2:
            st.caption("Bootstrap admins no se pueden desactivar.")

# =========================
# MODULE: COMMUNITIES (ADMIN)
# =========================
if module == "🏢 Comunidades":
    st.header("🏢 Módulo Comunidades (Admin)")

    st.subheader("Crear comunidad")
    new_name = st.text_input("Nombre comunidad", placeholder="Ej: Edificio A - Los Castaños")
    if st.button("Crear"):
        if not new_name.strip():
            st.error("Falta nombre.")
        else:
            row = create_community(new_name.strip())
            # Admin se da acceso automáticamente
            grant_access(auth.email, row["community_id"], "admin", True, True)
            st.success(f"Creada: {row['community_name']} ({row['community_id']})")
            st.rerun()

    st.divider()
    st.subheader("Gestionar comunidad")

    comms_all = list_communities(include_inactive=True)
    if not comms_all:
        st.info("No hay comunidades.")
        st.stop()

    comm_labels = [f"{c['community_name']} ({c['community_id']})" for c in comms_all]
    idx = st.selectbox("Comunidad", options=list(range(len(comms_all))), format_func=lambda i: comm_labels[i], index=0)
    comm = comms_all[idx]
    cid = comm["community_id"]
    cname = comm["community_name"]
    is_active = safe_bool_str(comm.get("is_active", ""), True) == "TRUE"

    colA, colB, colC = st.columns([2, 1, 2])
    with colA:
        st.write(f"**{cname}**  \n`{cid}`")
    with colB:
        if st.button("Desactivar" if is_active else "Reactivar"):
            set_community_active(cid, not is_active)
            st.rerun()
    with colC:
        st.caption("Foto de comunidad (opcional, Drive)")
        up = st.file_uploader("Portada", type=["jpg", "jpeg", "png"], label_visibility="collapsed", key=f"cover_{cid}")
        if up is not None:
            if not FOTOS_ROOT_FOLDER_ID:
                st.error("Falta FOTOS_ROOT_FOLDER_ID en secrets.")
            else:
                if st.button("Guardar portada", key=f"save_cover_{cid}"):
                    comm_folder = get_community_folder_id(cid, cname)
                    cover_folder = drive_get_or_create_folder(comm_folder, "_cover")
                    ext = up.type.split("/")[-1] if up.type and "/" in up.type else "jpg"
                    file_id, web_link = drive_upload_image_bytes(
                        parent_folder_id=cover_folder,
                        filename=f"cover.{ext}",
                        content_type=up.type or "image/jpeg",
                        file_bytes=up.getvalue(),
                    )
                    update_community_cover(cid, file_id, web_link)
                    st.success("Portada guardada ✅")
                    st.rerun()

    st.divider()

    # Installations + tasks editor with "Guardar cambios"
    st.subheader("Instalaciones y Tareas (configuración por comunidad)")
    st.caption("Edita en modo borrador local y luego presiona **Guardar cambios**.")

    # Load current catalog
    inst_all = list_installations(cid, include_inactive=True)
    tasks_all = list_tasks(cid, include_inactive=True)

    if "draft_catalog" not in st.session_state or st.session_state.get("draft_catalog_cid") != cid:
        # prepare draft structures
        st.session_state["draft_catalog_cid"] = cid
        st.session_state["draft_catalog"] = {
            "installations": [dict(x) for x in inst_all],
            "tasks": [dict(x) for x in tasks_all],
        }

    draft = st.session_state["draft_catalog"]
    draft_insts = draft["installations"]
    draft_tasks = draft["tasks"]

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Precargar plantilla sugerida (si está vacío)"):
            if inst_all or tasks_all:
                st.warning("Ya hay datos. Si quieres limpiar, usa Reset DB (testing) en la barra lateral.")
            else:
                seed_default_installations_and_tasks_if_empty(cid)
                st.success("Plantilla precargada.")
                st.session_state.pop("draft_catalog", None)
                st.rerun()

    with col2:
        st.info("Tip: desactiva (no borres) para mantener históricos coherentes.")

    st.markdown("### Instalaciones")
    with st.expander("➕ Agregar instalación"):
        cat = st.selectbox("Categoría", options=DEFAULT_CATEGORIES + ["Otra"], index=0, key="cat_add")
        cat_other = ""
        if cat == "Otra":
            cat_other = st.text_input("Nombre categoría", placeholder="Ej: Seguridad")
        name = st.text_input("Nombre instalación", placeholder="Ej: Portón norte")
        if st.button("Agregar instalación"):
            category = cat_other.strip() if cat == "Otra" else cat
            if not category or not name.strip():
                st.error("Completa categoría y nombre.")
            else:
                # add to draft only
                draft_insts.append({
                    "installation_id": f"INS-DRAFT-{secrets.token_hex(4)}",
                    "community_id": cid,
                    "category": category,
                    "installation_name": name.strip(),
                    "is_active": "TRUE",
                    "created_at": iso(now_utc()),
                    "_draft_new": "TRUE"
                })
                st.success("Agregada al borrador. No olvides Guardar cambios.")

    # Show draft installations
    for ins in sorted(draft_insts, key=lambda x: (x.get("category",""), x.get("installation_name",""))):
        active = safe_bool_str(ins.get("is_active",""), True) == "TRUE"
        with st.container(border=True):
            cA, cB, cC, cD = st.columns([2, 2, 1, 1])
            with cA:
                ins["installation_name"] = st.text_input("Instalación", value=ins.get("installation_name",""), key=f"ins_name_{ins['installation_id']}")
            with cB:
                ins["category"] = st.text_input("Categoría", value=ins.get("category",""), key=f"ins_cat_{ins['installation_id']}")
            with cC:
                ins["is_active"] = "TRUE" if st.checkbox("Activa", value=active, key=f"ins_act_{ins['installation_id']}") else "FALSE"
            with cD:
                st.caption(f"ID: {ins['installation_id']}")

            # Tasks for installation
            st.markdown("**Tareas**")
            related = [t for t in draft_tasks if t.get("installation_id") == ins["installation_id"]]
            related.sort(key=lambda x: int(float(x.get("task_order","9999") or 9999)) if str(x.get("task_order","")).strip() else 9999)

            for t in related:
                t_active = safe_bool_str(t.get("is_active",""), True) == "TRUE"
                t_req = safe_bool_str(t.get("is_required",""), True) == "TRUE"
                cc1, cc2, cc3, cc4 = st.columns([3, 1, 1, 1])
                with cc1:
                    t["task_name"] = st.text_input("Tarea", value=t.get("task_name",""), key=f"tsk_name_{t['task_id']}")
                with cc2:
                    t["task_order"] = st.text_input("Orden", value=str(t.get("task_order","10")), key=f"tsk_ord_{t['task_id']}")
                with cc3:
                    t["is_required"] = "TRUE" if st.checkbox("Req.", value=t_req, key=f"tsk_req_{t['task_id']}") else "FALSE"
                with cc4:
                    t["is_active"] = "TRUE" if st.checkbox("Activa", value=t_active, key=f"tsk_act_{t['task_id']}") else "FALSE"

            with st.expander("➕ Agregar tarea a esta instalación"):
                tn = st.text_input("Nombre tarea", key=f"new_tn_{ins['installation_id']}", placeholder="Ej: Revisión de fugas")
                to = st.number_input("Orden", min_value=1, value=10, step=10, key=f"new_to_{ins['installation_id']}")
                tr = st.checkbox("Es requerida", value=True, key=f"new_tr_{ins['installation_id']}")
                if st.button("Agregar tarea", key=f"btn_add_task_{ins['installation_id']}"):
                    if not tn.strip():
                        st.error("Falta nombre tarea.")
                    else:
                        draft_tasks.append({
                            "task_id": f"TSK-DRAFT-{secrets.token_hex(4)}",
                            "community_id": cid,
                            "installation_id": ins["installation_id"],
                            "task_name": tn.strip(),
                            "task_order": str(int(to)),
                            "is_required": "TRUE" if tr else "FALSE",
                            "is_active": "TRUE",
                            "created_at": iso(now_utc()),
                            "_draft_new": "TRUE"
                        })
                        st.success("Tarea agregada al borrador. Guardar cambios para persistir.")

    st.divider()
    if st.button("💾 Guardar cambios (persistir en Sheets)"):
        # Persist draft changes:
        # - New draft IDs are replaced with real IDs and updated accordingly for tasks.
        # - Existing IDs update by rewriting whole tables for community subset.
        try:
            # Load real tables
            inst_headers, inst_rows = read_table("Installations")
            t_headers, t_rows = read_table("Tasks")

            # Remove old rows of this community (we will re-add from draft)
            inst_rows = [r for r in inst_rows if r.get("community_id") != cid]
            t_rows = [r for r in t_rows if r.get("community_id") != cid]

            # Map draft installation IDs to real IDs
            inst_id_map = {}
            for ins in draft_insts:
                did = ins["installation_id"]
                if did.startswith("INS-DRAFT-"):
                    rid = f"INS-{secrets.token_hex(4)}"
                    inst_id_map[did] = rid
                else:
                    inst_id_map[did] = did

            # Write installations for this community
            for ins in draft_insts:
                real_id = inst_id_map[ins["installation_id"]]
                inst_rows.append({
                    "installation_id": real_id,
                    "community_id": cid,
                    "category": ins.get("category","").strip(),
                    "installation_name": ins.get("installation_name","").strip(),
                    "is_active": safe_bool_str(ins.get("is_active","TRUE"), True),
                    "created_at": ins.get("created_at","") or iso(now_utc()),
                })

            # Tasks: map installation_id using inst_id_map; map draft task ids to real
            for t in draft_tasks:
                tid = t["task_id"]
                if tid.startswith("TSK-DRAFT-"):
                    real_tid = f"TSK-{secrets.token_hex(4)}"
                else:
                    real_tid = tid

                t_rows.append({
                    "task_id": real_tid,
                    "community_id": cid,
                    "installation_id": inst_id_map.get(t.get("installation_id",""), t.get("installation_id","")),
                    "task_name": (t.get("task_name","") or "").strip(),
                    "task_order": str(int(float(t.get("task_order","10") or 10))),
                    "is_required": safe_bool_str(t.get("is_required","TRUE"), True),
                    "is_active": safe_bool_str(t.get("is_active","TRUE"), True),
                    "created_at": t.get("created_at","") or iso(now_utc()),
                })

            write_table("Installations", TABS_HEADERS["Installations"], inst_rows)
            write_table("Tasks", TABS_HEADERS["Tasks"], t_rows)

            st.success("Cambios guardados ✅")
            # refresh draft from DB
            st.session_state.pop("draft_catalog", None)
            st.rerun()
        except Exception as e:
            st.error("Error guardando cambios.")
            st.exception(e)

# =========================
# MODULE: REPORTS
# =========================
if module == "🧾 Informes":
    st.header("🧾 Módulo Informes")

    if not allowed_comms:
        st.warning("No tienes comunidades asignadas.")
        st.stop()

    labels = [f"{x['community_name']} ({x['community_id']})" for x in allowed_comms]
    sel_i = st.selectbox("Comunidad", options=list(range(len(allowed_comms))), format_func=lambda i: labels[i], index=0)
    comm = allowed_comms[sel_i]
    cid = comm["community_id"]
    cname = comm["community_name"]
    role = comm["role"]
    can_create = comm["can_create_reports"] or auth.is_admin
    if not can_create:
        st.warning("No tienes permiso para crear/editar informes en esta comunidad.")
        st.stop()

    # Ensure catalog exists
    inst = list_installations(cid, include_inactive=False)
    tasks = list_tasks(cid, include_inactive=False)
    if not inst or not tasks:
        st.warning("Esta comunidad no tiene catálogo completo (Instalaciones y/o Tareas). Pide al admin configurarlo.")
        st.stop()

    st.caption("Elige: crear nuevo o continuar Draft (último) para esta comunidad.")

    colA, colB = st.columns([1, 1])
    with colA:
        report_date = st.date_input("Fecha del informe", value=date.today())
    with colB:
        draft_window_hours = INVITE_EXPIRY_HOURS  # reutilizamos 48 por defecto, si quieres lo separamos luego
        st.caption(f"Continuar Draft busca el último Draft del usuario en esta comunidad (ventana ~{draft_window_hours}h).")

    c1, c2 = st.columns(2)
    with c1:
        if st.button("➕ Crear nuevo informe (Draft)", use_container_width=True):
            rep = create_new_report(cid, cname, auth.email, report_date)
            ensure_report_items_from_catalog(rep, auth.email)
            st.session_state["selected_report_id"] = rep["report_id"]
            st.success("Informe creado.")
            st.rerun()
    with c2:
        if st.button("⏩ Continuar informe pendiente (Draft)", use_container_width=True):
            reps = list_reports_for_user_and_community(auth.email, cid)
            cutoff = now_utc() - timedelta(hours=draft_window_hours)
            pick = None
            for r in reps:
                if r.get("status") != "Draft":
                    continue
                upd = parse_iso(r.get("updated_at","")) or parse_iso(r.get("created_at",""))
                if upd and upd >= cutoff:
                    pick = r
                    break
            if not pick:
                st.warning("No encontré un Draft reciente para continuar.")
            else:
                ensure_report_items_from_catalog(pick, auth.email)
                st.session_state["selected_report_id"] = pick["report_id"]
                st.success(f"Cargando Draft: {pick['report_id']}")
                st.rerun()

    st.divider()

    report_id = st.session_state.get("selected_report_id", "")
    if report_id:
        report = get_report(report_id)
        if not report:
            st.warning("No se encontró el informe seleccionado.")
            st.session_state.pop("selected_report_id", None)
            st.stop()

        # Ensure items exist
        ensure_report_items_from_catalog(report, auth.email)
        items = list_report_items(report_id)

        st.subheader(f"Informe: {report_id}")
        st.write(f"**Comunidad:** {report.get('community_name','')} (`{report.get('community_id','')}`)")

        # Stats cards
        ok_n = sum(1 for it in items if (it.get("status","") or "").lower() == "ok")
        fail_n = sum(1 for it in items if (it.get("status","") or "").lower() == "fail")
        pend_n = sum(1 for it in items if (it.get("status","") or "").lower() == "pending")

        st.markdown(
            f"""
            <div style="display:flex; gap:12px; flex-wrap:wrap; margin:8px 0 8px 0;">
              <div style="padding:12px 14px; border-radius:14px; border:1px solid #e2e8f0; min-width:160px;">
                <div style="font-size:12px; opacity:0.7;">OK</div>
                <div style="font-size:22px; font-weight:800; color:#16a34a;">{ok_n}</div>
              </div>
              <div style="padding:12px 14px; border-radius:14px; border:1px solid #e2e8f0; min-width:160px;">
                <div style="font-size:12px; opacity:0.7;">FALLAS</div>
                <div style="font-size:22px; font-weight:800; color:#dc2626;">{fail_n}</div>
              </div>
              <div style="padding:12px 14px; border-radius:14px; border:1px solid #e2e8f0; min-width:160px;">
                <div style="font-size:12px; opacity:0.7;">PENDIENTES</div>
                <div style="font-size:22px; font-weight:800; color:#64748b;">{pend_n}</div>
              </div>
            </div>
            """,
            unsafe_allow_html=True
        )

        # Status + save
        col1, col2, col3 = st.columns([1, 1, 2])
        with col1:
            new_status = st.selectbox("Estado del informe", options=REPORT_STATUS, index=REPORT_STATUS.index(report.get("status","Draft")))
        with col2:
            if st.button("Guardar estado"):
                set_report_status(report_id, new_status)
                st.success("Estado guardado.")
                st.rerun()
        with col3:
            st.caption("Los cambios de checklist se guardan al vuelo por cada ítem/tarea.")

        # Checklist UI grouped by category / installation
        items_map = get_report_items_map(report_id)

        # sort categories
        cats = sorted(
            list(set([it.get("category","") for it in items])),
            key=lambda s: (DEFAULT_CATEGORIES.index(s) if s in DEFAULT_CATEGORIES else 999, s)
        )

        for cat in cats:
            st.markdown(f"## {cat}")
            inst_names = sorted(list(set([it.get("installation_name","") for it in items if it.get("category","") == cat])))
            for ins_name in inst_names:
                # Collect items for this installation
                ins_items = [it for it in items if it.get("category","") == cat and it.get("installation_name","") == ins_name]
                # Need installation_id for folder routing
                ins_id = ins_items[0].get("installation_id","")

                st.markdown(f"### {ins_name}")

                for it in ins_items:
                    status = (it.get("status","pending") or "pending").lower()
                    bg = status_bg(status)
                    with st.container(border=True):
                        st.markdown(
                            f"<div style='padding:8px 10px; border-radius:10px; background:{bg};'>"
                            f"<div style='font-weight:800; font-size:15px;'>{it.get('task_name','')}</div>"
                            f"</div>",
                            unsafe_allow_html=True
                        )

                        a, b, c = st.columns([1.2, 2.6, 2.2], gap="medium")
                        with a:
                            sel = st.radio(
                                "Estado",
                                options=STATUS_VALUES,
                                index=STATUS_VALUES.index(status if status in STATUS_VALUES else "pending"),
                                format_func=lambda v: {"pending":"Pendiente","ok":"OK","fail":"Falla"}[v],
                                horizontal=True,
                                label_visibility="collapsed",
                                key=f"st_{report_id}_{it['installation_id']}_{it['task_id']}",
                            )
                        with b:
                            note = st.text_input(
                                "Observación",
                                value=it.get("note",""),
                                placeholder="Observación breve…",
                                label_visibility="collapsed",
                                key=f"nt_{report_id}_{it['installation_id']}_{it['task_id']}",
                            )
                        with c:
                            st.caption("📷 Foto (Drive)")
                            if it.get("photo_web_view_link"):
                                st.markdown(f"✅ Vinculada: [Abrir]({it['photo_web_view_link']})")

                            up = st.file_uploader(
                                "Subir",
                                type=["jpg", "jpeg", "png"],
                                label_visibility="collapsed",
                                key=f"ph_{report_id}_{it['installation_id']}_{it['task_id']}",
                            )
                            if up is not None:
                                if not FOTOS_ROOT_FOLDER_ID:
                                    st.error("Falta FOTOS_ROOT_FOLDER_ID en secrets.")
                                else:
                                    if st.button("Subir y vincular", key=f"upl_{report_id}_{it['installation_id']}_{it['task_id']}"):
                                        try:
                                            comm_folder = get_community_folder_id(cid, cname)
                                            rep_folder = get_report_folder_id(comm_folder, report_id)
                                            ins_folder = get_installation_folder_id(rep_folder, ins_id, ins_name)
                                            ext = up.type.split("/")[-1] if up.type and "/" in up.type else "jpg"
                                            filename = f"{report_id}__{it['installation_id']}__{it['task_id']}.{ext}"
                                            file_id, web_link = drive_upload_image_bytes(
                                                parent_folder_id=ins_folder,
                                                filename=filename,
                                                content_type=up.type or "image/jpeg",
                                                file_bytes=up.getvalue(),
                                            )
                                            it["photo_file_id"] = file_id
                                            it["photo_web_view_link"] = web_link
                                            st.success("Foto vinculada ✅")
                                        except Exception as e:
                                            st.error("Error subiendo foto.")
                                            st.exception(e)

                        # Persist updates (status/note/photo) immediately
                        it["status"] = sel
                        it["note"] = note
                        it["updated_at"] = iso(now_utc())
                        it["updated_by"] = norm_email(auth.email)
                        upsert_report_item(it)

        st.divider()
        st.subheader("Descargar informe")
        st.caption("Descarga con estructura visual (categoría / instalación / tareas) e incluye fotos vinculadas, si existen.")

        colx, coly = st.columns(2)
        with colx:
            if REPORTLAB_OK:
                if st.button("Generar PDF"):
                    try:
                        pdf_bytes = export_pdf(get_report(report_id))
                        st.download_button(
                            "⬇️ Descargar PDF",
                            data=pdf_bytes,
                            file_name=f"{report_id}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error("No pude generar el PDF.")
                        st.exception(e)
            else:
                st.info("PDF: instala reportlab para habilitar descarga.")

        with coly:
            if PYDOCX_OK:
                if st.button("Generar Word (.docx)"):
                    try:
                        docx_bytes = export_docx(get_report(report_id))
                        st.download_button(
                            "⬇️ Descargar Word",
                            data=docx_bytes,
                            file_name=f"{report_id}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error("No pude generar el Word.")
                        st.exception(e)
            else:
                st.info("Word: instala python-docx para habilitar descarga.")

# =========================
# MODULE: SUMMARY
# =========================
if module == "📊 Resumen":
    st.header("📊 Módulo Resumen")

    # All reports
    _, reports = read_table("Reports")
    reports = [r for r in reports if r.get("report_id")]

    # Filter by allowed communities if not admin
    allowed_ids = {c["community_id"] for c in allowed_comms}
    if not auth.is_admin:
        reports = [r for r in reports if r.get("community_id") in allowed_ids]

    # Filters UI
    communities_filter = sorted(list(set([r.get("community_name","") for r in reports if r.get("community_name")])))

    col1, col2, col3, col4 = st.columns([2, 2, 2, 2])
    with col1:
        f_comm = st.multiselect("Comunidad", options=communities_filter)
    with col2:
        f_user = st.text_input("Filtrar por email creador (contiene)", placeholder="ej: juan@")
    with col3:
        f_state = st.multiselect("Estado", options=REPORT_STATUS)
    with col4:
        f_date = st.date_input("Rango fecha (inicio/fin)", value=(date.today() - timedelta(days=30), date.today()))

    start_d, end_d = f_date if isinstance(f_date, tuple) else (f_date, f_date)

    def in_range(dstr: str) -> bool:
        try:
            d = date.fromisoformat(dstr)
            return start_d <= d <= end_d
        except Exception:
            return True

    filtered = []
    for r in reports:
        if f_comm and r.get("community_name","") not in f_comm:
            continue
        if f_user and f_user.lower() not in (r.get("created_by","") or "").lower():
            continue
        if f_state and r.get("status","") not in f_state:
            continue
        if r.get("report_date") and not in_range(r["report_date"]):
            continue
        filtered.append(r)

    filtered.sort(key=lambda x: (x.get("report_date",""), x.get("updated_at","")), reverse=True)

    st.write(f"Resultados: **{len(filtered)}**")

    # Simple table
    st.dataframe(
        [{
            "Comunidad": r.get("community_name",""),
            "Fecha": r.get("report_date",""),
            "Usuario": r.get("created_by",""),
            "Estado": r.get("status",""),
            "Report ID": r.get("report_id",""),
        } for r in filtered],
        use_container_width=True,
        hide_index=True
    )

    st.divider()
    st.subheader("Descargar un informe desde el historial")

    if not filtered:
        st.info("No hay informes para descargar con los filtros actuales.")
    else:
        pick = st.selectbox("Selecciona Report ID", options=[r["report_id"] for r in filtered])
        rep = get_report(pick)
        if rep:
            colx, coly = st.columns(2)
            with colx:
                if REPORTLAB_OK and st.button("Generar PDF", key="sum_pdf"):
                    try:
                        pdf_bytes = export_pdf(rep)
                        st.download_button(
                            "⬇️ Descargar PDF",
                            data=pdf_bytes,
                            file_name=f"{pick}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error("No pude generar el PDF.")
                        st.exception(e)
            with coly:
                if PYDOCX_OK and st.button("Generar Word", key="sum_docx"):
                    try:
                        docx_bytes = export_docx(rep)
                        st.download_button(
                            "⬇️ Descargar Word",
                            data=docx_bytes,
                            file_name=f"{pick}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    except Exception as e:
                        st.error("No pude generar el Word.")
                        st.exception(e)

st.caption("Control Comunidades • Esquema nuevo con Tasks/Reports/ReportItems • Testing OK")
