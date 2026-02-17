# app.py
# ============================================================
# 🛡️ Control Comunidades — Informes con Sheets + Export PDF/Word
#    - Sheets (Service Account) = Base de datos
#    - Drive (OAuth) = Guardar PDFs/DOCX en "Mi unidad" del admin
# ============================================================

import base64
import hashlib
import io
import json
import os
import re
import time
import uuid
from dataclasses import dataclass
from datetime import datetime, timedelta, date
from typing import Dict, List, Optional, Tuple

import pandas as pd
import streamlit as st
from PIL import Image

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError

from google_auth_oauthlib.flow import Flow
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request

# Report exports
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.styles import ParagraphStyle

from docx import Document
from docx.shared import Inches


# ============================================================
# CONFIG
# ============================================================

st.set_page_config(page_title="Control Comunidades", page_icon="🛡️", layout="wide")

APP_NAME = "🛡️ Control Comunidades"
NOW = lambda: datetime.now()

# ---------------- Secrets (MODO B) ----------------
def _sec(key: str, default=None):
    return st.secrets.get(key, default)

GCP_TYPE = _sec("GCP_TYPE")
GCP_PROJECT_ID = _sec("GCP_PROJECT_ID")
GCP_PRIVATE_KEY_ID = _sec("GCP_PRIVATE_KEY_ID")
GCP_PRIVATE_KEY = _sec("GCP_PRIVATE_KEY")
GCP_CLIENT_EMAIL = _sec("GCP_CLIENT_EMAIL")
GCP_CLIENT_ID = _sec("GCP_CLIENT_ID")
GCP_AUTH_URI = _sec("GCP_AUTH_URI")
GCP_TOKEN_URI = _sec("GCP_TOKEN_URI")
GCP_AUTH_PROVIDER_X509_CERT_URL = _sec("GCP_AUTH_PROVIDER_X509_CERT_URL")
GCP_CLIENT_X509_CERT_URL = _sec("GCP_CLIENT_X509_CERT_URL")

SHEET_ID = _sec("SHEET_ID")
BOOTSTRAP_ADMIN_EMAILS = _sec("BOOTSTRAP_ADMIN_EMAILS", "")
APP_PEPPER = _sec("APP_PEPPER", "pepper")
INVITE_EXPIRY_HOURS = int(_sec("INVITE_EXPIRY_HOURS", "48"))
BOOTSTRAP_SETUP_TOKEN = _sec("BOOTSTRAP_SETUP_TOKEN", "")
FOTOS_ROOT_FOLDER_ID = _sec("FOTOS_ROOT_FOLDER_ID", "")  # opcional (no se usa si incrustas foto en informe)
OAUTH_CLIENT_ID = _sec("OAUTH_CLIENT_ID", "")
OAUTH_CLIENT_SECRET = _sec("OAUTH_CLIENT_SECRET", "")
OAUTH_REDIRECT_URI = _sec("OAUTH_REDIRECT_URI", "")
DRIVE_REPORTS_FOLDER_ID = _sec("DRIVE_REPORTS_FOLDER_ID", "")

BOOTSTRAP_ADMINS = [e.strip().lower() for e in BOOTSTRAP_ADMIN_EMAILS.split(",") if e.strip()]

# Drive OAuth scopes: mínimo para subir/crear archivos
DRIVE_SCOPES = ["https://www.googleapis.com/auth/drive.file"]

# Roles
ROLE_ADMIN = "Admin"
ROLE_EDITOR = "Editor"   # crea informes + ve resumen
ROLE_VIEWER = "Viewer"   # solo ve resumen
NON_ADMIN_ROLES = ["Supervisor", "Mantenedor", "Analista Interno", "Conserje"]  # etiqueta/UX
ROLE_MAP = {
    "Supervisor": ROLE_EDITOR,
    "Mantenedor": ROLE_EDITOR,
    "Analista Interno": ROLE_EDITOR,
    "Conserje": ROLE_EDITOR,  # puedes cambiar a Viewer si quieres
}

# Sheet tabs
TAB_USERS = "Users"
TAB_INVITES = "Invites"
TAB_COMMUNITIES = "Communities"
TAB_INSTALLATIONS = "Installations"
TAB_TASKS = "Tasks"
TAB_REPORTS = "Reports"
TAB_REPORTITEMS = "ReportItems"

USERS_HEADERS = [
    "user_id", "email", "role", "role_label", "status",
    "communities_csv",
    "pw_salt", "pw_hash",
    "created_at", "last_login",
]
INVITES_HEADERS = [
    "invite_id", "email", "role", "role_label", "communities_csv",
    "code_salt", "code_hash",
    "expires_at", "created_at", "created_by", "used_at",
]
COMMUNITIES_HEADERS = [
    "community_id", "name", "is_active", "photo_drive_file_id",
    "created_at", "updated_at",
]
INSTALLATIONS_HEADERS = [
    "installation_id", "community_id", "category", "installation",
    "created_at", "updated_at",
]
TASKS_HEADERS = [
    "task_id", "installation_id", "task",
    "created_at", "updated_at",
]
REPORTS_HEADERS = [
    "report_id", "community_id", "report_date", "status",
    "created_by", "created_at", "updated_at",
    "drive_pdf_file_id", "drive_pdf_link",
    "drive_docx_file_id", "drive_docx_link",
]
REPORTITEMS_HEADERS = [
    "report_item_id", "report_id", "category", "installation", "task",
    "status", "note",
    "photo_mime", "photo_b64",
    "updated_at",
]

CATEGORIES_DEFAULT = ["Críticos", "Infraestructura", "Espacios Comunes", "Accesos", "Higiene", "Otros"]


# ============================================================
# UTIL: HASHING
# ============================================================

def _rand_salt() -> str:
    return base64.urlsafe_b64encode(os.urandom(16)).decode("utf-8").rstrip("=")

def _hash_with_pepper(value: str, salt: str) -> str:
    # PBKDF2-ish simple: (salt + value + pepper) hashed multiple times
    data = (salt + value + APP_PEPPER).encode("utf-8")
    h = hashlib.sha256(data).digest()
    for _ in range(60_000):
        h = hashlib.sha256(h).digest()
    return base64.urlsafe_b64encode(h).decode("utf-8").rstrip("=")

def _norm_email(email: str) -> str:
    return (email or "").strip().lower()

def _now_iso() -> str:
    return NOW().strftime("%Y-%m-%d %H:%M:%S")

def _today_iso() -> str:
    return date.today().isoformat()


# ============================================================
# GOOGLE SHEETS (Service Account)
# ============================================================

def _service_account_info() -> dict:
    # Construye dict compatible con google.oauth2.service_account
    return {
        "type": GCP_TYPE,
        "project_id": GCP_PROJECT_ID,
        "private_key_id": GCP_PRIVATE_KEY_ID,
        "private_key": (GCP_PRIVATE_KEY or "").replace("\\n", "\n"),
        "client_email": GCP_CLIENT_EMAIL,
        "client_id": GCP_CLIENT_ID,
        "auth_uri": GCP_AUTH_URI,
        "token_uri": GCP_TOKEN_URI,
        "auth_provider_x509_cert_url": GCP_AUTH_PROVIDER_X509_CERT_URL,
        "client_x509_cert_url": GCP_CLIENT_X509_CERT_URL,
    }

@st.cache_resource(show_spinner=False)
def sheets_service():
    info = _service_account_info()
    creds = service_account.Credentials.from_service_account_info(
        info,
        scopes=["https://www.googleapis.com/auth/spreadsheets"],
    )
    return build("sheets", "v4", credentials=creds)

def _sheets_call(fn, *args, **kwargs):
    # Retry simple para 429/503
    backoff = 0.8
    for attempt in range(6):
        try:
            return fn(*args, **kwargs)
        except HttpError as e:
            status = getattr(e, "status_code", None) or getattr(getattr(e, "resp", None), "status", None)
            if status in (429, 500, 503):
                time.sleep(backoff)
                backoff *= 1.8
                continue
            raise

def sheets_get_spreadsheet(spreadsheet_id: str):
    svc = sheets_service()
    return _sheets_call(svc.spreadsheets().get(spreadsheetId=spreadsheet_id).execute)

def sheets_values_get(range_a1: str):
    svc = sheets_service()
    return _sheets_call(
        svc.spreadsheets().values().get(spreadsheetId=SHEET_ID, range=range_a1).execute
    )

def sheets_values_update(range_a1: str, values: List[List[str]]):
    svc = sheets_service()
    body = {"values": values}
    return _sheets_call(
        svc.spreadsheets().values().update(
            spreadsheetId=SHEET_ID,
            range=range_a1,
            valueInputOption="RAW",
            body=body,
        ).execute
    )

def sheets_values_append(range_a1: str, values: List[List[str]]):
    svc = sheets_service()
    body = {"values": values}
    return _sheets_call(
        svc.spreadsheets().values().append(
            spreadsheetId=SHEET_ID,
            range=range_a1,
            valueInputOption="RAW",
            insertDataOption="INSERT_ROWS",
            body=body,
        ).execute
    )

def sheets_batch_update(requests: list):
    svc = sheets_service()
    body = {"requests": requests}
    return _sheets_call(
        svc.spreadsheets().batchUpdate(spreadsheetId=SHEET_ID, body=body).execute
    )

def _ensure_sheet_tab(tab_name: str):
    meta = sheets_get_spreadsheet(SHEET_ID)
    tabs = {s["properties"]["title"] for s in meta.get("sheets", [])}
    if tab_name in tabs:
        return
    sheets_batch_update([{
        "addSheet": {"properties": {"title": tab_name}}
    }])

def _ensure_headers(tab: str, headers: List[str]):
    _ensure_sheet_tab(tab)
    a1 = f"{tab}!A1:Z1"
    got = sheets_values_get(a1).get("values", [])
    if got and got[0] == headers:
        return
    sheets_values_update(f"{tab}!A1", [headers])

def ensure_schema():
    if st.session_state.get("_schema_ok"):
        return
    for t, h in [
        (TAB_USERS, USERS_HEADERS),
        (TAB_INVITES, INVITES_HEADERS),
        (TAB_COMMUNITIES, COMMUNITIES_HEADERS),
        (TAB_INSTALLATIONS, INSTALLATIONS_HEADERS),
        (TAB_TASKS, TASKS_HEADERS),
        (TAB_REPORTS, REPORTS_HEADERS),
        (TAB_REPORTITEMS, REPORTITEMS_HEADERS),
    ]:
        _ensure_headers(t, h)
    st.session_state["_schema_ok"] = True

def _sheet_to_df(tab: str, headers: List[str]) -> pd.DataFrame:
    ensure_schema()
    data = sheets_values_get(f"{tab}!A1:Z").get("values", [])
    if not data:
        return pd.DataFrame(columns=headers)
    got_headers = data[0]
    rows = data[1:]
    # Normaliza a largo de headers esperado
    col_count = len(headers)
    normalized = []
    for r in rows:
        r = (r + [""] * col_count)[:col_count]
        normalized.append(r)
    df = pd.DataFrame(normalized, columns=headers)
    return df

def _df_to_sheet_overwrite(tab: str, headers: List[str], df: pd.DataFrame):
    ensure_schema()
    df2 = df.copy()
    # asegura columnas
    for c in headers:
        if c not in df2.columns:
            df2[c] = ""
    df2 = df2[headers]
    values = [headers] + df2.astype(str).fillna("").values.tolist()
    sheets_values_update(f"{tab}!A1", values)

def _append_row(tab: str, headers: List[str], row: dict):
    ensure_schema()
    values = [[str(row.get(h, "")) for h in headers]]
    sheets_values_append(f"{tab}!A1", values)

def _update_row_by_id(tab: str, headers: List[str], id_col: str, row_id: str, updates: dict):
    df = _sheet_to_df(tab, headers)
    if df.empty:
        return False
    mask = df[id_col].astype(str) == str(row_id)
    if not mask.any():
        return False
    for k, v in updates.items():
        if k in df.columns:
            df.loc[mask, k] = str(v)
    _df_to_sheet_overwrite(tab, headers, df)
    return True


# ============================================================
# GOOGLE DRIVE (OAuth) — subir PDFs/DOCX a "Mi unidad"
# ============================================================

def _oauth_client_config():
    return {
        "web": {
            "client_id": OAUTH_CLIENT_ID,
            "client_secret": OAUTH_CLIENT_SECRET,
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "redirect_uris": [OAUTH_REDIRECT_URI],
        }
    }

def _get_drive_creds_from_state() -> Optional[Credentials]:
    raw = st.session_state.get("drive_oauth_token")
    if not raw:
        return None
    creds = Credentials.from_authorized_user_info(raw, scopes=DRIVE_SCOPES)
    if creds.expired and creds.refresh_token:
        try:
            creds.refresh(Request())
            st.session_state["drive_oauth_token"] = json.loads(creds.to_json())
        except Exception:
            return None
    return creds

@st.cache_resource(show_spinner=False)
def _drive_service_cached(token_json_str: str):
    raw = json.loads(token_json_str)
    creds = Credentials.from_authorized_user_info(raw, scopes=DRIVE_SCOPES)
    return build("drive", "v3", credentials=creds)

def drive_service() -> Optional[object]:
    creds = _get_drive_creds_from_state()
    if not creds:
        return None
    return _drive_service_cached(creds.to_json())

def drive_upload_bytes(
    file_bytes: bytes,
    filename: str,
    mime_type: str,
    parent_folder_id: str,
) -> Tuple[str, str]:
    svc = drive_service()
    if svc is None:
        raise RuntimeError("Drive no está conectado (OAuth). Conéctalo con un admin.")
    from googleapiclient.http import MediaIoBaseUpload

    media = MediaIoBaseUpload(io.BytesIO(file_bytes), mimetype=mime_type, resumable=False)
    body = {"name": filename}
    if parent_folder_id:
        body["parents"] = [parent_folder_id]
    created = _sheets_call(
        svc.files().create(
            body=body,
            media_body=media,
            fields="id, webViewLink",
        ).execute
    )
    return created["id"], created.get("webViewLink", "")


# ============================================================
# BOOTSTRAP + AUTH
# ============================================================

def get_users_df() -> pd.DataFrame:
    return _sheet_to_df(TAB_USERS, USERS_HEADERS)

def get_invites_df() -> pd.DataFrame:
    return _sheet_to_df(TAB_INVITES, INVITES_HEADERS)

def get_communities_df() -> pd.DataFrame:
    df = _sheet_to_df(TAB_COMMUNITIES, COMMUNITIES_HEADERS)
    if not df.empty:
        df["is_active"] = df["is_active"].astype(str).str.lower().isin(["true", "1", "yes", "y"])
    return df

def get_installations_df() -> pd.DataFrame:
    return _sheet_to_df(TAB_INSTALLATIONS, INSTALLATIONS_HEADERS)

def get_tasks_df() -> pd.DataFrame:
    return _sheet_to_df(TAB_TASKS, TASKS_HEADERS)

def get_reports_df() -> pd.DataFrame:
    return _sheet_to_df(TAB_REPORTS, REPORTS_HEADERS)

def get_reportitems_df() -> pd.DataFrame:
    return _sheet_to_df(TAB_REPORTITEMS, REPORTITEMS_HEADERS)

def _bootstrap_needed() -> bool:
    df = get_users_df()
    return df.empty

def _bootstrap_create_admins(setup_token: str):
    if not BOOTSTRAP_SETUP_TOKEN:
        raise RuntimeError("Falta BOOTSTRAP_SETUP_TOKEN en secrets.")
    if setup_token.strip() != BOOTSTRAP_SETUP_TOKEN.strip():
        raise RuntimeError("BOOTSTRAP_SETUP_TOKEN inválido.")

    now = _now_iso()
    for email in BOOTSTRAP_ADMINS:
        user_id = str(uuid.uuid4())
        pw_salt = _rand_salt()
        # password inicial: se fuerza a setear luego (guardamos hash vacío)
        row = {
            "user_id": user_id,
            "email": email,
            "role": ROLE_ADMIN,
            "role_label": ROLE_ADMIN,
            "status": "active",
            "communities_csv": "*",
            "pw_salt": pw_salt,
            "pw_hash": "",  # sin pass, se setea al entrar
            "created_at": now,
            "last_login": "",
        }
        _append_row(TAB_USERS, USERS_HEADERS, row)

def _user_has_password(user_row: pd.Series) -> bool:
    return bool(str(user_row.get("pw_hash", "")).strip())


def _set_user_password(email: str, new_password: str):
    """Setea (o resetea) la contraseña del/los registros del email.
    Si existen duplicados, los actualiza TODOS para evitar inconsistencias.
    """
    df = get_users_df()
    em = _norm_email(email)
    mask = df["email"].astype(str).str.lower() == em
    if not mask.any():
        raise RuntimeError("Usuario no existe.")

    sub = df.loc[mask].copy()
    salt_candidates = sub.get("pw_salt", "").astype(str).tolist() if "pw_salt" in sub.columns else []
    salt = next((s for s in reversed(salt_candidates) if s and s.strip()), "") or _rand_salt()

    pw_hash = _hash_with_pepper(new_password, salt)
    df.loc[mask, "pw_salt"] = salt
    df.loc[mask, "pw_hash"] = pw_hash
    if "updated_at" in df.columns:
        df.loc[mask, "updated_at"] = _now_iso()
    _df_to_sheet_overwrite(TAB_USERS, USERS_HEADERS, df)



def _verify_password(email: str, password: str) -> Optional[pd.Series]:
    """Valida credenciales. Si hay duplicados, prioriza:
    - status=active
    - registro con pw_hash no vacío
    - el más reciente (por updated_at/created_at)
    """
    df = get_users_df()
    em = _norm_email(email)
    rows = df[df["email"].astype(str).str.lower() == em].copy()
    if rows.empty:
        return None

    rows["__status"] = rows.get("status", "").astype(str).str.lower()
    rows = rows[rows["__status"].fillna("") == "active"].copy()
    if rows.empty:
        return None

    # ordena por tiempo si existe
    time_col = "updated_at" if "updated_at" in rows.columns else ("created_at" if "created_at" in rows.columns else None)
    if time_col:
        rows["__t"] = rows[time_col].astype(str)
        rows = rows.sort_values("__t")

    # prioriza los que tienen contraseña
    rows["__has_pw"] = rows.get("pw_hash", "").astype(str).str.strip().ne("")
    # Queremos el mejor al final: primero sin pw, luego con pw
    rows = rows.sort_values(["__has_pw"])

    r = rows.iloc[-1]

    if not _user_has_password(r):
        return r  # se pedirá setear

    salt = str(r.get("pw_salt", ""))
    expected = str(r.get("pw_hash", ""))
    got = _hash_with_pepper(password, salt)
    return r if got == expected else None


def _can_access_community(user_row: dict, community_id: str) -> bool:
    csv = (user_row.get("communities_csv") or "").strip()
    if csv == "*" and user_row.get("role") == ROLE_ADMIN:
        return True
    allowed = [x.strip() for x in csv.split(",") if x.strip()]
    return community_id in allowed or "*" in allowed

def _role_base(role: str) -> str:
    if role == ROLE_ADMIN:
        return ROLE_ADMIN
    if role in (ROLE_EDITOR, ROLE_VIEWER):
        return role
    return ROLE_MAP.get(role, ROLE_EDITOR)

def _login_user(user_row: dict):
    st.session_state["auth_user"] = {
        "user_id": user_row.get("user_id"),
        "email": _norm_email(user_row.get("email")),
        "role": user_row.get("role"),
        "role_label": user_row.get("role_label") or user_row.get("role"),
        "communities_csv": user_row.get("communities_csv", ""),
    }
    _update_row_by_id(
        TAB_USERS, USERS_HEADERS, "user_id", user_row["user_id"],
        {"last_login": _now_iso()}
    )

def _logout():
    st.session_state.pop("auth_user", None)
    st.session_state.pop("drive_oauth_token", None)
    st.query_params.clear()
    st.rerun()

def current_user() -> Optional[dict]:
    return st.session_state.get("auth_user")


# ============================================================
# INVITES (Admin)
# ============================================================

def create_invite(email: str, role: str, role_label: str, communities_csv: str, created_by: str) -> str:
    email = _norm_email(email)
    if not email:
        raise RuntimeError("Email inválido.")
    invite_id = str(uuid.uuid4())
    code_plain = uuid.uuid4().hex[:10].upper()  # código humano
    code_salt = _rand_salt()
    code_hash = _hash_with_pepper(code_plain, code_salt)

    expires_at = (NOW() + timedelta(hours=INVITE_EXPIRY_HOURS)).strftime("%Y-%m-%d %H:%M:%S")
    row = {
        "invite_id": invite_id,
        "email": email,
        "role": role,
        "role_label": role_label,
        "communities_csv": communities_csv,
        "code_salt": code_salt,
        "code_hash": code_hash,
        "expires_at": expires_at,
        "created_at": _now_iso(),
        "created_by": created_by,
        "used_at": "",
    }
    _append_row(TAB_INVITES, INVITES_HEADERS, row)
    return code_plain

def redeem_invite(email: str, code_plain: str) -> Tuple[bool, str]:
    email = _norm_email(email)
    code_plain = (code_plain or "").strip().upper()
    df = get_invites_df()
    if df.empty:
        return False, "No hay invitaciones."

    # filtra por email y no usadas
    df2 = df[
        (df["email"].astype(str).str.lower() == email) &
        (df["used_at"].astype(str).str.strip() == "")
    ].copy()

    if df2.empty:
        return False, "No existe una invitación activa para ese email."

    # valida cada una (podrían existir varias)
    for _, inv in df2.iterrows():
        try:
            exp = datetime.strptime(inv["expires_at"], "%Y-%m-%d %H:%M:%S")
        except Exception:
            continue
        if NOW() > exp:
            continue

        salt = str(inv.get("code_salt", ""))
        expected = str(inv.get("code_hash", ""))
        got = _hash_with_pepper(code_plain, salt)
        if got != expected:
            continue

        # crea usuario si no existe
        users = get_users_df()
        exists = not users[users["email"].astype(str).str.lower() == email].empty
        if not exists:
            user_id = str(uuid.uuid4())
            pw_salt = _rand_salt()
            row = {
                "user_id": user_id,
                "email": email,
                "role": inv["role"],
                "role_label": inv["role_label"],
                "status": "active",
                "communities_csv": inv["communities_csv"],
                "pw_salt": pw_salt,
                "pw_hash": "",
                "created_at": _now_iso(),
                "last_login": "",
            }
            _append_row(TAB_USERS, USERS_HEADERS, row)

        # marca invite usada
        _update_row_by_id(
            TAB_INVITES, INVITES_HEADERS, "invite_id", inv["invite_id"],
            {"used_at": _now_iso()}
        )
        return True, "Invitación aceptada. Ahora define tu contraseña."

    return False, "Código inválido o expirado."


# ============================================================
# COMMUNITIES / INSTALLATIONS / TASKS (Admin)
# ============================================================

def create_community(name: str) -> str:
    name = (name or "").strip()
    if not name:
        raise RuntimeError("Nombre inválido.")
    comm_id = str(uuid.uuid4())
    now = _now_iso()
    row = {
        "community_id": comm_id,
        "name": name,
        "is_active": "TRUE",
        "photo_drive_file_id": "",
        "created_at": now,
        "updated_at": now,
    }
    _append_row(TAB_COMMUNITIES, COMMUNITIES_HEADERS, row)
    return comm_id

def delete_community_permanent(community_id: str):
    # Borra: comunidad + instalaciones + tareas + reportes + reportitems
    community_id = str(community_id)
    cdf = get_communities_df()
    cdf = cdf[cdf["community_id"].astype(str) != community_id].copy()
    _df_to_sheet_overwrite(TAB_COMMUNITIES, COMMUNITIES_HEADERS, cdf)

    idf = get_installations_df()
    to_del_inst = idf[idf["community_id"].astype(str) == community_id]["installation_id"].astype(str).tolist()
    idf = idf[idf["community_id"].astype(str) != community_id].copy()
    _df_to_sheet_overwrite(TAB_INSTALLATIONS, INSTALLATIONS_HEADERS, idf)

    tdf = get_tasks_df()
    if to_del_inst:
        tdf = tdf[~tdf["installation_id"].astype(str).isin(to_del_inst)].copy()
        _df_to_sheet_overwrite(TAB_TASKS, TASKS_HEADERS, tdf)

    rdf = get_reports_df()
    to_del_reports = rdf[rdf["community_id"].astype(str) == community_id]["report_id"].astype(str).tolist()
    rdf = rdf[rdf["community_id"].astype(str) != community_id].copy()
    _df_to_sheet_overwrite(TAB_REPORTS, REPORTS_HEADERS, rdf)

    ridf = get_reportitems_df()
    if to_del_reports:
        ridf = ridf[~ridf["report_id"].astype(str).isin(to_del_reports)].copy()
        _df_to_sheet_overwrite(TAB_REPORTITEMS, REPORTITEMS_HEADERS, ridf)

def upsert_installation(community_id: str, category: str, installation: str) -> str:
    community_id = str(community_id)
    category = (category or "").strip() or "Otros"
    installation = (installation or "").strip()
    if not installation:
        raise RuntimeError("Instalación inválida.")

    df = get_installations_df()
    now = _now_iso()

    # dedupe natural: (community_id, category, installation)
    if not df.empty:
        mask = (
            (df["community_id"].astype(str) == community_id) &
            (df["category"].astype(str) == category) &
            (df["installation"].astype(str) == installation)
        )
        if mask.any():
            inst_id = df.loc[mask, "installation_id"].astype(str).iloc[0]
            _update_row_by_id(TAB_INSTALLATIONS, INSTALLATIONS_HEADERS, "installation_id", inst_id, {"updated_at": now})
            return inst_id

    inst_id = str(uuid.uuid4())
    row = {
        "installation_id": inst_id,
        "community_id": community_id,
        "category": category,
        "installation": installation,
        "created_at": now,
        "updated_at": now,
    }
    _append_row(TAB_INSTALLATIONS, INSTALLATIONS_HEADERS, row)
    return inst_id

def delete_installation(installation_id: str):
    installation_id = str(installation_id)
    idf = get_installations_df()
    idf = idf[idf["installation_id"].astype(str) != installation_id].copy()
    _df_to_sheet_overwrite(TAB_INSTALLATIONS, INSTALLATIONS_HEADERS, idf)

    tdf = get_tasks_df()
    tdf = tdf[tdf["installation_id"].astype(str) != installation_id].copy()
    _df_to_sheet_overwrite(TAB_TASKS, TASKS_HEADERS, tdf)

def upsert_task(installation_id: str, task: str) -> str:
    installation_id = str(installation_id)
    task = (task or "").strip()
    if not task:
        raise RuntimeError("Tarea inválida.")
    df = get_tasks_df()
    now = _now_iso()

    if not df.empty:
        mask = (df["installation_id"].astype(str) == installation_id) & (df["task"].astype(str) == task)
        if mask.any():
            tid = df.loc[mask, "task_id"].astype(str).iloc[0]
            _update_row_by_id(TAB_TASKS, TASKS_HEADERS, "task_id", tid, {"updated_at": now})
            return tid

    tid = str(uuid.uuid4())
    row = {
        "task_id": tid,
        "installation_id": installation_id,
        "task": task,
        "created_at": now,
        "updated_at": now,
    }
    _append_row(TAB_TASKS, TASKS_HEADERS, row)
    return tid

def delete_task(task_id: str):
    task_id = str(task_id)
    tdf = get_tasks_df()
    tdf = tdf[tdf["task_id"].astype(str) != task_id].copy()
    _df_to_sheet_overwrite(TAB_TASKS, TASKS_HEADERS, tdf)


# ============================================================
# REPORTS (Editor/Admin)
# ============================================================

def create_report(community_id: str, report_date: str, created_by: str) -> str:
    rid = str(uuid.uuid4())
    now = _now_iso()
    row = {
        "report_id": rid,
        "community_id": str(community_id),
        "report_date": report_date,
        "status": "draft",
        "created_by": created_by,
        "created_at": now,
        "updated_at": now,
        "drive_pdf_file_id": "",
        "drive_pdf_link": "",
        "drive_docx_file_id": "",
        "drive_docx_link": "",
    }
    _append_row(TAB_REPORTS, REPORTS_HEADERS, row)
    return rid

def set_report_status(report_id: str, status: str):
    status = status.lower().strip()
    if status not in ("draft", "final"):
        status = "draft"
    _update_row_by_id(TAB_REPORTS, REPORTS_HEADERS, "report_id", report_id, {"status": status, "updated_at": _now_iso()})

def get_latest_draft_report(community_id: str) -> Optional[str]:
    df = get_reports_df()
    if df.empty:
        return None
    df2 = df[
        (df["community_id"].astype(str) == str(community_id)) &
        (df["status"].astype(str).str.lower() == "draft")
    ].copy()
    if df2.empty:
        return None
    # orden por updated_at
    def parse_dt(x):
        try:
            return datetime.strptime(str(x), "%Y-%m-%d %H:%M:%S")
        except Exception:
            return datetime(1970, 1, 1)
    df2["__dt"] = df2["updated_at"].apply(parse_dt)
    df2 = df2.sort_values("__dt", ascending=False)
    latest = df2.iloc[0]
    # expiración
    if parse_dt(latest["updated_at"]) < NOW() - timedelta(hours=INVITE_EXPIRY_HOURS):
        return None
    return str(latest["report_id"])

def build_report_matrix(community_id: str) -> pd.DataFrame:
    # Devuelve filas por (category, installation, task)
    idf = get_installations_df()
    tdf = get_tasks_df()

    idf = idf[idf["community_id"].astype(str) == str(community_id)].copy()
    if idf.empty:
        return pd.DataFrame(columns=["category", "installation", "task", "installation_id", "task_id"])

    tdf = tdf.copy()
    rows = []
    for _, inst in idf.iterrows():
        inst_id = inst["installation_id"]
        cat = inst["category"]
        inst_name = inst["installation"]
        tasks = tdf[tdf["installation_id"].astype(str) == str(inst_id)].copy()
        if tasks.empty:
            rows.append({
                "category": cat,
                "installation": inst_name,
                "task": "(sin tareas)",
                "installation_id": str(inst_id),
                "task_id": "",
            })
        else:
            for _, t in tasks.iterrows():
                rows.append({
                    "category": cat,
                    "installation": inst_name,
                    "task": t["task"],
                    "installation_id": str(inst_id),
                    "task_id": str(t["task_id"]),
                })
    out = pd.DataFrame(rows)
    # orden básico
    out["__cat"] = out["category"].astype(str)
    out = out.sort_values(["__cat", "installation", "task"]).drop(columns=["__cat"])
    return out

def load_reportitems(report_id: str) -> pd.DataFrame:
    df = get_reportitems_df()
    if df.empty:
        return pd.DataFrame(columns=REPORTITEMS_HEADERS)
    return df[df["report_id"].astype(str) == str(report_id)].copy()

def save_reportitems(report_id: str, items_df: pd.DataFrame):
    # Sobrescribe todas las filas del report_id (simple y robusto para testing)
    all_df = get_reportitems_df()
    other = all_df[all_df["report_id"].astype(str) != str(report_id)].copy()
    combined = pd.concat([other, items_df], ignore_index=True)
    _df_to_sheet_overwrite(TAB_REPORTITEMS, REPORTITEMS_HEADERS, combined)
    _update_row_by_id(TAB_REPORTS, REPORTS_HEADERS, "report_id", report_id, {"updated_at": _now_iso()})


# ============================================================
# EXPORT: PDF (ReportLab) / WORD (python-docx)
# ============================================================

def _status_label(s: str) -> str:
    s = (s or "").lower().strip()
    return {"ok": "OK", "fail": "FALLA", "pending": "PEND."}.get(s, "PEND.")

def _status_color_bg(s: str):
    s = (s or "").lower().strip()
    if s == "fail":
        return colors.Color(1, 0.93, 0.93)  # rojo suave
    if s == "ok":
        return colors.Color(0.93, 1, 0.93)  # verde suave
    return colors.whitesmoke

def export_pdf_visual(
    community_name: str,
    report_date_str: str,
    rows: List[dict],
) -> bytes:
    # rows: [{category, installation, task, status, note, photo_bytes|None, photo_mime}]
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "title",
        parent=styles["Heading1"],
        fontSize=16,
        leading=20,
        alignment=TA_LEFT,
    )
    meta_style = ParagraphStyle(
        "meta",
        parent=styles["Normal"],
        fontSize=10,
        leading=13,
    )

    flow = []
    flow.append(Paragraph(f"{APP_NAME} — Informe", title_style))
    flow.append(Paragraph(f"<b>Comunidad:</b> {community_name}", meta_style))
    flow.append(Paragraph(f"<b>Fecha informe:</b> {report_date_str}", meta_style))
    flow.append(Spacer(1, 8))

    # Header table
    data = [["Categoría / Instalación / Tarea", "Estado / Observación", "Foto"]]

    # Build cells
    for r in rows:
        left = f"<b>{r['category']}</b><br/>{r['installation']}<br/><i>{r['task']}</i>"
        mid = f"<b>{_status_label(r.get('status'))}</b><br/>{(r.get('note') or '').strip() or '—'}"
        img_cell = "—"
        if r.get("photo_bytes"):
            try:
                img = Image.open(io.BytesIO(r["photo_bytes"]))
                img = img.convert("RGB")
                img.thumbnail((240, 240))
                ib = io.BytesIO()
                img.save(ib, format="JPEG", quality=85)
                ib.seek(0)
                img_cell = RLImage(ib, width=55 * mm, height=55 * mm)  # cuadrado
            except Exception:
                img_cell = "Foto inválida"

        data.append([Paragraph(left, meta_style), Paragraph(mid, meta_style), img_cell])

    table = Table(data, colWidths=[70 * mm, 60 * mm, 55 * mm])
    ts = TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 10),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fafafa")]),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ])

    # Fallas con fondo rojo suave
    for i, r in enumerate(rows, start=1):
        bg = _status_color_bg(r.get("status"))
        ts.add("BACKGROUND", (0, i), (-1, i), bg)

    table.setStyle(ts)
    flow.append(table)

    doc.build(flow)
    return buf.getvalue()

def export_docx_visual(
    community_name: str,
    report_date_str: str,
    rows: List[dict],
) -> bytes:
    doc = Document()
    doc.add_heading(f"{APP_NAME} — Informe", level=1)
    doc.add_paragraph(f"Comunidad: {community_name}")
    doc.add_paragraph(f"Fecha informe: {report_date_str}")

    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Categoría / Instalación / Tarea"
    hdr[1].text = "Estado / Observación"
    hdr[2].text = "Foto"

    for r in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{r['category']}\n{r['installation']}\n{r['task']}"
        row_cells[1].text = f"{_status_label(r.get('status'))}\n{(r.get('note') or '').strip() or '—'}"
        if r.get("photo_bytes"):
            try:
                img = Image.open(io.BytesIO(r["photo_bytes"]))
                img = img.convert("RGB")
                img.thumbnail((600, 600))
                ib = io.BytesIO()
                img.save(ib, format="JPEG", quality=85)
                ib.seek(0)
                run = row_cells[2].paragraphs[0].add_run()
                run.add_picture(ib, width=Inches(1.8))
            except Exception:
                row_cells[2].text = "Foto inválida"
        else:
            row_cells[2].text = "—"

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ============================================================
# MASTER DATA TEMPLATE (xlsx)
# ============================================================

def build_master_template_xlsx() -> bytes:
    # Plantilla: Community | Category | Installation | Task
    from openpyxl import Workbook
    from openpyxl.worksheet.datavalidation import DataValidation

    wb = Workbook()
    ws = wb.active
    ws.title = "MasterData"
    ws.append(["CommunityName", "Category", "Installation", "Task"])
    ws.append(["Ej: Comunidad A", "Críticos", "Sala de Bombas", "Presión y alternancia"])

    # Data validation for Category
    dv = DataValidation(type="list", formula1='"{}"'.format(",".join(CATEGORIES_DEFAULT)), allow_blank=False)
    ws.add_data_validation(dv)
    dv.add("B2:B500")

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()

def import_master_template_xlsx(file_bytes: bytes, target_community_id: str, replace_all: bool):
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    if "MasterData" not in wb.sheetnames:
        raise RuntimeError("La plantilla debe tener una hoja llamada 'MasterData'.")
    ws = wb["MasterData"]

    rows = []
    for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        community_name, category, installation, task = row
        if not community_name and not category and not installation and not task:
            continue
        category = (category or "Otros").strip()
        installation = (installation or "").strip()
        task = (task or "").strip()
        if not installation:
            continue
        rows.append((category, installation, task))

    if replace_all:
        # elimina instalaciones + tareas de esa comunidad y vuelve a crear
        idf = get_installations_df()
        inst_ids = idf[idf["community_id"].astype(str) == str(target_community_id)]["installation_id"].astype(str).tolist()
        idf = idf[idf["community_id"].astype(str) != str(target_community_id)].copy()
        _df_to_sheet_overwrite(TAB_INSTALLATIONS, INSTALLATIONS_HEADERS, idf)

        tdf = get_tasks_df()
        if inst_ids:
            tdf = tdf[~tdf["installation_id"].astype(str).isin(inst_ids)].copy()
            _df_to_sheet_overwrite(TAB_TASKS, TASKS_HEADERS, tdf)

    # upsert
    for category, installation, task in rows:
        inst_id = upsert_installation(target_community_id, category, installation)
        if task:
            upsert_task(inst_id, task)


# ============================================================
# OAUTH UI (Drive)
# ============================================================

def drive_oauth_ui(user: dict):
    st.sidebar.markdown("### 🧩 Drive (Mi unidad)")
    creds = _get_drive_creds_from_state()
    if creds:
        st.sidebar.success("Drive conectado (OAuth).")
        if st.sidebar.button("Desconectar Drive", use_container_width=True):
            st.session_state.pop("drive_oauth_token", None)
            st.query_params.clear()
            st.rerun()
        return

    if user.get("role") != ROLE_ADMIN:
        st.sidebar.info("Solo Admin conecta Drive.")
        return

    st.sidebar.warning("Drive NO conectado.")

    # Handle redirect callback
    qp = st.query_params
    code = qp.get("code")
    if code:
        try:
            flow = Flow.from_client_config(
                _oauth_client_config(),
                scopes=DRIVE_SCOPES,
                redirect_uri=OAUTH_REDIRECT_URI,
            )
            flow.fetch_token(code=code)
            creds = flow.credentials
            st.session_state["drive_oauth_token"] = json.loads(creds.to_json())
            st.query_params.clear()
            st.sidebar.success("Drive conectado ✅")
            st.rerun()
        except Exception as e:
            st.sidebar.error(f"Error OAuth: {e}")
            st.query_params.clear()

    if not (OAUTH_CLIENT_ID and OAUTH_CLIENT_SECRET and OAUTH_REDIRECT_URI):
        st.sidebar.error("Faltan secrets OAuth (OAUTH_CLIENT_ID/SECRET/REDIRECT_URI).")
        return

    if st.sidebar.button("Conectar Drive (OAuth)", use_container_width=True):
        flow = Flow.from_client_config(
            _oauth_client_config(),
            scopes=DRIVE_SCOPES,
            redirect_uri=OAUTH_REDIRECT_URI,
        )
        auth_url, _ = flow.authorization_url(
            access_type="offline",
            include_granted_scopes="true",
            prompt="consent",
        )
        st.sidebar.markdown("Abre este enlace para autorizar:")
        st.sidebar.markdown(auth_url)


# ============================================================
# UI: HEADER
# ============================================================

def header_card(ok: int, fail: int, pending: int):
    st.markdown(
        f"""
        <div style="padding:18px 18px 12px 18px; border-radius:16px;
                    background: linear-gradient(90deg, #111827, #1f2937);
                    color:white;">
          <div style="display:flex; justify-content:space-between; align-items:center; gap:16px; flex-wrap:wrap;">
            <div>
              <div style="font-size:22px; font-weight:900;">{APP_NAME}</div>
              <div style="opacity:0.85;">Sheets (DB) + Drive (export) • esquema Tasks/Reports</div>
            </div>
            <div style="display:flex; gap:14px;">
              <div style="background: rgba(255,255,255,0.08); padding:10px 14px; border-radius:12px; text-align:center; min-width:110px;">
                <div style="font-size:11px; letter-spacing:0.08em; opacity:0.8;">OK</div>
                <div style="font-size:22px; font-weight:900; color:#34d399;">{ok}</div>
              </div>
              <div style="background: rgba(255,255,255,0.08); padding:10px 14px; border-radius:12px; text-align:center; min-width:110px;">
                <div style="font-size:11px; letter-spacing:0.08em; opacity:0.8;">FALLAS</div>
                <div style="font-size:22px; font-weight:900; color:#fb7185;">{fail}</div>
              </div>
              <div style="background: rgba(255,255,255,0.08); padding:10px 14px; border-radius:12px; text-align:center; min-width:110px;">
                <div style="font-size:11px; letter-spacing:0.08em; opacity:0.8;">PEND.</div>
                <div style="font-size:22px; font-weight:900; color:#fbbf24;">{pending}</div>
              </div>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


# ============================================================
# UI: LOGIN SCREEN
# ============================================================


def login_screen():
    st.title(f"{APP_NAME} — Acceso")
    st.caption("Acceso seguro por invitación (Admin) + contraseña. Incluye modo de recuperación (Administrador Principal / Bootstrap).")

    ensure_schema()

    # =========================
    # 1) Ingreso con contraseña (DEFAULT)
    # =========================
    st.subheader("✅ Ingreso con Contraseña")
    with st.form("login_form"):
        email = st.text_input("Email", placeholder="usuario@empresa.com")
        password = st.text_input("Contraseña", type="password")
        submit = st.form_submit_button("Entrar")

    if submit:
        user = _verify_password(email, password)
        if user is None:
            st.error("Credenciales inválidas.")
        else:
            # Si no tiene password aún, forzamos setear
            if not _user_has_password(user):
                st.session_state["pw_setup_email"] = _norm_email(email)
                st.success("Primer ingreso: define una contraseña.")
            else:
                _login_user(user.to_dict())
                st.rerun()

    if "pw_setup_email" in st.session_state:
        st.info(f"Define contraseña para: {st.session_state['pw_setup_email']}")
        with st.form("set_pw_form"):
            p1 = st.text_input("Nueva contraseña", type="password")
            p2 = st.text_input("Repite contraseña", type="password")
            ok = st.form_submit_button("Guardar contraseña")
        if ok:
            if len(p1) < 8:
                st.error("Usa al menos 8 caracteres.")
            elif p1 != p2:
                st.error("No coinciden.")
            else:
                try:
                    _set_user_password(st.session_state["pw_setup_email"], p1)
                    # auto-login tras set password
                    user2 = _verify_password(st.session_state["pw_setup_email"], p1)
                    if (user2 is not None) and _user_has_password(user2):
                        _login_user(user2.to_dict())
                        st.session_state.pop("pw_setup_email", None)
                        st.success("Contraseña definida. Ingreso exitoso.")
                        st.rerun()
                    else:
                        st.warning("Contraseña guardada pero no pudo verificarse. Revisa APP_PEPPER / duplicados en Users.")
                except Exception as e:
                    st.error(str(e))

    st.divider()

    # =========================
    # 2) Ingreso con código de un uso (Plegado)
    # =========================
    with st.expander("🔑 Ingreso con Código de un Uso (Invitación)", expanded=False):
        st.caption("Si un Admin te dio un código, actívalo aquí y luego define tu contraseña.")
        with st.form("redeem_form"):
            email2 = st.text_input("Email (invitación)", key="inv_email")
            code = st.text_input("Código", key="inv_code")
            submit2 = st.form_submit_button("Activar")
        if submit2:
            ok, msg = redeem_invite(email2, code)
            if ok:
                st.success(msg)
                st.session_state["pw_setup_email"] = _norm_email(email2)
            else:
                st.error(msg)

    # =========================
    # 3) Administrador Principal (Bootstrap / Recuperación) (Plegado)
    # =========================
    with st.expander("🧰 Ingreso Administrador Principal (Bootstrap / Recuperación)", expanded=False):
        st.caption("Usa esta opción para primer ingreso o recuperación. Requiere token + email bootstrap autorizado.")
        if not BOOTSTRAP_SETUP_TOKEN:
            st.error("Falta BOOTSTRAP_SETUP_TOKEN en secrets.")
        else:
            users_df = get_users_df()
            st.info("Acciones disponibles: (A) Crear admins iniciales si la base está vacía; (B) Definir/Resetear contraseña de un bootstrap admin.")
            action = st.radio("Acción", ["Definir/Resetear contraseña (Bootstrap)", "Crear Admins iniciales (solo si no hay usuarios)"], index=0)

            if action.startswith("Crear"):
                with st.form("bootstrap_create_admins"):
                    token = st.text_input("BOOTSTRAP_SETUP_TOKEN", type="password", key="bs_create_token")
                    submit3 = st.form_submit_button("Crear Admins iniciales")
                if submit3:
                    try:
                        _bootstrap_create_admins(token)
                        st.success("Admins bootstrap creados. Ahora define contraseña usando la otra acción.")
                        st.rerun()
                    except Exception as e:
                        st.error(str(e))
            else:
                with st.form("bootstrap_reset_pw"):
                    token = st.text_input("BOOTSTRAP_SETUP_TOKEN", type="password", key="bs_token")
                    emailb = st.text_input("Email Admin Principal", placeholder=(BOOTSTRAP_ADMINS[0] if BOOTSTRAP_ADMINS else "admin@empresa.com"), key="bs_email").strip().lower()
                    p1 = st.text_input("Nueva contraseña", type="password", key="bs_p1")
                    p2 = st.text_input("Repite contraseña", type="password", key="bs_p2")
                    submit4 = st.form_submit_button("Guardar y entrar")

                if submit4:
                    try:
                        if token.strip() != BOOTSTRAP_SETUP_TOKEN.strip():
                            raise RuntimeError("BOOTSTRAP_SETUP_TOKEN inválido.")
                        if emailb not in BOOTSTRAP_ADMINS:
                            raise RuntimeError("Email no autorizado como bootstrap admin.")
                        if len(p1) < 8:
                            raise RuntimeError("Usa al menos 8 caracteres.")
                        if p1 != p2:
                            raise RuntimeError("No coinciden.")

                        # si no existe el usuario, lo creamos como admin total
                        dfu = get_users_df()
                        exists = not dfu[dfu["email"].astype(str).str.lower() == emailb].empty
                        if not exists:
                            now = _now_iso()
                            row = {
                                "user_id": str(uuid.uuid4()),
                                "email": emailb,
                                "role": ROLE_ADMIN,
                                "role_label": ROLE_ADMIN,
                                "status": "active",
                                "communities_csv": "*",
                                "pw_salt": _rand_salt(),
                                "pw_hash": "",
                                "created_at": now,
                                "last_login": "",
                            }
                            _append_row(TAB_USERS, USERS_HEADERS, row)

                        _set_user_password(emailb, p1)

                        user = _verify_password(emailb, p1)
                        if user is None:
                            raise RuntimeError("Contraseña guardada pero no pudo verificarse (revisa APP_PEPPER/duplicados).")
                        _login_user(user.to_dict())
                        st.success("Ingreso exitoso (Bootstrap).")
                        st.rerun()
                    except Exception as e:
                        st.error(str(e))


# ============================================================
# UI: MODULES

# ============================================================

def module_users(user: dict):
    st.subheader("👥 Módulo Usuarios (Admins)")
    st.caption("Admins pueden crear invitaciones, asignar roles y comunidades, y eliminar usuarios (excepto bootstrap).")

    users = get_users_df()
    invites = get_invites_df()
    comms = get_communities_df()

    st.markdown("#### Crear invitación")
    with st.form("invite_form"):
        email = st.text_input("Email a invitar")
        role_label = st.selectbox("Rol (etiqueta)", [ROLE_ADMIN] + NON_ADMIN_ROLES + [ROLE_VIEWER], index=1)
        base_role = ROLE_ADMIN if role_label == ROLE_ADMIN else (ROLE_VIEWER if role_label == ROLE_VIEWER else ROLE_MAP.get(role_label, ROLE_EDITOR))

        # Comunidades a asignar
        comm_options = [(r["community_id"], r["name"]) for _, r in comms[comms["is_active"]].iterrows()] if not comms.empty else []
        selected = st.multiselect("Comunidades (vacío = ninguna)", options=comm_options, format_func=lambda x: x[1])
        communities_csv = ",".join([cid for (cid, _) in selected]) if selected else ""
        if base_role == ROLE_ADMIN:
            communities_csv = "*"  # Admin acceso total

        submit = st.form_submit_button("Generar código")
    if submit:
        try:
            code = create_invite(
                email=email,
                role=base_role,
                role_label=role_label,
                communities_csv=communities_csv,
                created_by=user["email"],
            )
            st.success("Invitación creada ✅")
            st.code(f"CÓDIGO: {code}", language="text")
            st.info("Comparte el código al usuario por un canal seguro.")
        except Exception as e:
            st.error(str(e))

    st.divider()

    st.markdown("#### Usuarios activos")
    if users.empty:
        st.info("Sin usuarios.")
        return

    show = users.copy()
    show["is_bootstrap"] = show["email"].astype(str).str.lower().isin(BOOTSTRAP_ADMINS)
    st.dataframe(show[["email", "role", "role_label", "status", "communities_csv", "last_login", "is_bootstrap"]], use_container_width=True)

    st.markdown("#### Editar / Eliminar usuario")
    emails = show["email"].astype(str).tolist()
    pick = st.selectbox("Selecciona usuario", emails)
    row = show[show["email"].astype(str) == pick].iloc[0].to_dict()

    c1, c2, c3 = st.columns([1, 1, 1.2])
    with c1:
        new_role_label = st.selectbox("Rol (etiqueta)", [ROLE_ADMIN] + NON_ADMIN_ROLES + [ROLE_VIEWER], index=([ROLE_ADMIN] + NON_ADMIN_ROLES + [ROLE_VIEWER]).index(row["role_label"]) if row["role_label"] in ([ROLE_ADMIN] + NON_ADMIN_ROLES + [ROLE_VIEWER]) else 1)
    with c2:
        new_base_role = ROLE_ADMIN if new_role_label == ROLE_ADMIN else (ROLE_VIEWER if new_role_label == ROLE_VIEWER else ROLE_MAP.get(new_role_label, ROLE_EDITOR))
        st.write(f"Rol técnico: **{new_base_role}**")
    with c3:
        new_status = st.selectbox("Estado", ["active", "disabled"], index=0 if row["status"] == "active" else 1)

    # Comunidades
    comms = get_communities_df()
    comm_options = [(r["community_id"], r["name"]) for _, r in comms[comms["is_active"]].iterrows()] if not comms.empty else []
    current_csv = (row.get("communities_csv") or "").strip()
    current_set = set([x.strip() for x in current_csv.split(",") if x.strip()])
    preselect = [opt for opt in comm_options if opt[0] in current_set]

    if new_base_role == ROLE_ADMIN:
        st.info("Admin: comunidades se fijan a '*' (acceso total).")
        chosen_csv = "*"
    else:
        chosen = st.multiselect("Comunidades asignadas", options=comm_options, default=preselect, format_func=lambda x: x[1])
        chosen_csv = ",".join([cid for (cid, _) in chosen])

    colS, colD = st.columns([1, 1])
    with colS:
        if st.button("💾 Guardar cambios", use_container_width=True):
            try:
                df = get_users_df()
                mask = df["email"].astype(str).str.lower() == _norm_email(pick)
                if not mask.any():
                    st.error("No encontrado.")
                else:
                    df.loc[mask, "role"] = new_base_role
                    df.loc[mask, "role_label"] = new_role_label
                    df.loc[mask, "status"] = new_status
                    df.loc[mask, "communities_csv"] = chosen_csv
                    _df_to_sheet_overwrite(TAB_USERS, USERS_HEADERS, df)
                    st.success("Actualizado ✅")
                    st.rerun()
            except Exception as e:
                st.error(str(e))

    with colD:
        is_bootstrap = _norm_email(pick) in BOOTSTRAP_ADMINS
        if is_bootstrap:
            st.warning("Bootstrap: no se puede eliminar.")
        else:
            if st.button("🗑️ Eliminar usuario definitivamente", use_container_width=True):
                try:
                    df = get_users_df()
                    df = df[df["email"].astype(str).str.lower() != _norm_email(pick)].copy()
                    _df_to_sheet_overwrite(TAB_USERS, USERS_HEADERS, df)
                    st.success("Eliminado ✅")
                    st.rerun()
                except Exception as e:
                    st.error(str(e))

    st.divider()
    st.markdown("#### Invitaciones pendientes")
    inv = invites.copy()
    if inv.empty:
        st.info("No hay invitaciones.")
    else:
        inv["is_active"] = inv["used_at"].astype(str).str.strip().eq("")
        st.dataframe(inv[["email", "role_label", "role", "communities_csv", "expires_at", "created_by", "used_at"]], use_container_width=True)


def module_communities(user: dict):
    st.subheader("🏢 Módulo Comunidades (Admins)")
    st.caption("Crear/Eliminar comunidades y administrar instalaciones + tareas por comunidad. (Foto de comunidad opcional)")

    # Siempre refrescamos desde Sheets en cada render del módulo
    comms = get_communities_df()

    # ---------------------------
    # 1) Crear / Eliminar
    # ---------------------------
    with st.expander("➕ Crear / 🗑️ Eliminar comunidades", expanded=True):
        st.markdown("#### Crear nueva comunidad")
        with st.form("create_comm", clear_on_submit=True):
            name = st.text_input("Nombre comunidad", placeholder="Ej: Edificio A / Comunidad Las Palmas")
            submitted = st.form_submit_button("Crear comunidad")
            if submitted:
                try:
                    comm_id = create_community(name)
                    st.success("Comunidad creada ✅")
                    st.code(comm_id)
                    # Limpia selección anterior para evitar inconsistencias
                    st.session_state.pop("comm_edit_pick", None)
                    st.rerun()
                except Exception as e:
                    st.error(str(e))

        st.divider()
        st.markdown("#### Comunidades existentes")
        if comms.empty:
            st.info("Aún no hay comunidades.")
        else:
            st.caption("Tip: si tienes nombres repetidos, fíjate en el ID corto.")
            # Tabla + acciones de borrado por fila (evita confusión del selectbox con nombres repetidos)
            for _, r in comms.sort_values(["name"]).iterrows():
                cid = str(r["community_id"])
                cname = str(r["name"])
                cactive = bool(r.get("is_active", True))
                left, mid, right = st.columns([4, 1.3, 1.7])
                with left:
                    st.markdown(f"**{cname}**  \n`{cid}`")
                with mid:
                    st.markdown("**Activa:** " + ("✅" if cactive else "⛔"))
                with right:
                    # Confirmación por comunidad (clave única)
                    confirm_key = f"del_confirm_{cid}"
                    confirm = st.text_input("Confirma", placeholder="ELIMINAR", key=confirm_key, label_visibility="collapsed")
                    if st.button("🗑️ Eliminar", key=f"del_btn_{cid}", disabled=(confirm.strip().upper() != "ELIMINAR")):
                        try:
                            delete_community_permanent(cid)
                            st.success(f"Eliminada ✅ ({cname})")
                            # Limpia estados de UI para que no quede una comunidad eliminada seleccionada
                            for k in ["comm_edit_pick", "del_comm_confirm"]:
                                st.session_state.pop(k, None)
                            st.session_state.pop(confirm_key, None)
                            st.rerun()
                        except Exception as e:
                            st.error(str(e))
                st.markdown("---")

    st.divider()

    # Re-carga por si se creó/eliminó algo
    comms = get_communities_df()
    if comms.empty:
        return

    # ---------------------------
    # 2) Modificar comunidades ya creadas
    # ---------------------------
    with st.expander("🛠️ Modificar comunidad (estado, foto, instalaciones y tareas)", expanded=True):
        # Opciones con ID corto para evitar confusión cuando hay nombres repetidos
        comm_opts = [(str(r["community_id"]), f"{str(r['name'])} — {str(r['community_id'])[:6]}") for _, r in comms.iterrows()]
        pick = st.selectbox("Selecciona comunidad", comm_opts, format_func=lambda x: x[1], key="comm_edit_pick")
        comm_id = pick[0]
        comm_name = pick[1].split(" — ")[0]

        # Row actual (robusto ante estados inconsistentes)
        row = comms[comms["community_id"].astype(str) == str(comm_id)]
        if row.empty:
            st.warning("La comunidad seleccionada ya no existe (probablemente fue eliminada). Refrescando…")
            st.session_state.pop("comm_edit_pick", None)
            st.rerun()
        is_active_now = bool(row["is_active"].iloc[0])

        colA, colB = st.columns([1, 1])
        with colA:
            st.markdown("#### Configuración comunidad")
            active = st.checkbox("Activa", value=is_active_now, key=f"active_{comm_id}")
            if st.button("Guardar estado", key=f"save_active_{comm_id}"):
                try:
                    _update_row_by_id(
                        TAB_COMMUNITIES,
                        COMMUNITIES_HEADERS,
                        "community_id",
                        comm_id,
                        {"is_active": "TRUE" if active else "FALSE", "updated_at": _now_iso()},
                    )
                    st.success("Actualizado ✅")
                    st.rerun()
                except Exception as e:
                    st.error(str(e))

            st.markdown("#### Foto (opcional)")
            st.caption("Para una interfaz más atractiva. (Se guarda en Drive solo si Drive está conectado)")
            up = st.file_uploader("Subir foto comunidad", type=["png", "jpg", "jpeg"], key=f"photo_comm_{comm_id}")
            if up and st.button("Guardar foto", key=f"save_photo_{comm_id}"):
                try:
                    img_bytes = up.getvalue()
                    fname = f"community_{comm_name}_{uuid.uuid4().hex[:6]}.jpg"
                    fid, link = drive_upload_bytes(img_bytes, fname, up.type or "image/jpeg", DRIVE_REPORTS_FOLDER_ID)
                    _update_row_by_id(
                        TAB_COMMUNITIES,
                        COMMUNITIES_HEADERS,
                        "community_id",
                        comm_id,
                        {"photo_drive_file_id": fid, "updated_at": _now_iso()},
                    )
                    st.success("Foto guardada ✅ (Drive)")
                except Exception as e:
                    st.error(str(e))

        with colB:
            st.markdown("#### Datos maestros (xlsx)")
            tpl = build_master_template_xlsx()
            st.download_button(
                "⬇️ Descargar plantilla MasterData.xlsx",
                data=tpl,
                file_name="MasterData.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key=f"dl_tpl_{comm_id}",
            )
            st.caption("Puedes cargar la plantilla para crear instalaciones + tareas. Opción de reemplazar todo.")
            upx = st.file_uploader("Cargar MasterData.xlsx", type=["xlsx"], key=f"master_upload_{comm_id}")
            replace = st.checkbox("Reemplazar TODO (instalaciones+tareas) de esta comunidad", value=False, key=f"replace_all_{comm_id}")
            if upx and st.button("Importar plantilla", key=f"import_tpl_{comm_id}"):
                try:
                    import_master_template_xlsx(upx.getvalue(), comm_id, replace_all=replace)
                    st.success("Importado ✅")
                    st.rerun()
                except Exception as e:
                    st.error(str(e))

        st.divider()
        st.markdown("## Instalaciones y tareas")

        idf = get_installations_df()
        tdf = get_tasks_df()
        idf2 = idf[idf["community_id"].astype(str) == str(comm_id)].copy()

        with st.expander("➕ Añadir instalación", expanded=False):
            cat = st.selectbox("Tipo/Categoría", CATEGORIES_DEFAULT, key=f"new_inst_cat_{comm_id}")
            inst = st.text_input("Nombre instalación", placeholder="Ej: Sala de Bombas", key=f"new_inst_name_{comm_id}")
            if st.button("Guardar instalación", key=f"save_inst_{comm_id}"):
                try:
                    inst_id = upsert_installation(comm_id, cat, inst)
                    st.success("Instalación guardada ✅")
                    st.code(inst_id)
                    st.rerun()
                except Exception as e:
                    st.error(str(e))

        if idf2.empty:
            st.info("No hay instalaciones para esta comunidad.")
            return

        # UI por instalación
        for _, inst_row in idf2.sort_values(["category", "installation"]).iterrows():
            inst_id = inst_row["installation_id"]
            cat = inst_row["category"]
            inst_name = inst_row["installation"]

            box = st.container(border=True)
            with box:
                h1, h2 = st.columns([3, 1])
                with h1:
                    st.markdown(f"**{cat} — {inst_name}**")
                with h2:
                    if st.button("Eliminar instalación", key=f"del_inst_{inst_id}"):
                        try:
                            delete_installation(inst_id)
                            st.success("Eliminada ✅")
                            st.rerun()
                        except Exception as e:
                            st.error(str(e))

                # tareas
                tasks = tdf[tdf["installation_id"].astype(str) == str(inst_id)].copy()
                if tasks.empty:
                    st.caption("Sin tareas (aún).")
                else:
                    for _, t in tasks.sort_values(["task"]).iterrows():
                        c1, c2 = st.columns([5, 1])
                        with c1:
                            st.write(f"- {t['task']}")
                        with c2:
                            if st.button("🗑️", key=f"del_task_{t['task_id']}"):
                                try:
                                    delete_task(t["task_id"])
                                    st.rerun()
                                except Exception as e:
                                    st.error(str(e))

                with st.expander("➕ Añadir tarea", expanded=False):
                    task_txt = st.text_input("Tarea", key=f"task_txt_{inst_id}", placeholder="Ej: Presión y alternancia")
                    if st.button("Guardar tarea", key=f"save_task_{inst_id}"):
                        try:
                            upsert_task(inst_id, task_txt)
                            st.success("Tarea guardada ✅")
                            st.rerun()
                        except Exception as e:
                            st.error(str(e))


def module_informes(user: dict):
    st.subheader("🧾 Módulo Informes (creación)")
    st.caption("Selecciona comunidad asignada, crea informe nuevo o continúa el borrador. Guarda Draft/Final y exporta PDF/Word (con fotos incrustadas).")

    comms = get_communities_df()
    if comms.empty:
        st.info("No hay comunidades.")
        return

    # comunidades visibles según permisos
    visible = []
    for _, r in comms[comms["is_active"]].iterrows():
        if user["role"] == ROLE_ADMIN:
            visible.append((r["community_id"], r["name"]))
        else:
            if _can_access_community(user, r["community_id"]):
                visible.append((r["community_id"], r["name"]))

    if not visible:
        st.warning("No tienes comunidades asignadas.")
        return

    pick = st.selectbox("Comunidad", visible, format_func=lambda x: x[1])
    community_id, community_name = pick[0], pick[1]

    # elegir informe
    c1, c2, c3 = st.columns([1, 1, 1.5])
    with c1:
        rep_date = st.date_input("Fecha informe", value=date.today())
        rep_date_str = rep_date.isoformat()
    with c2:
        latest_draft = get_latest_draft_report(community_id)
        action = st.radio(
            "Acción",
            options=["Nuevo informe", "Continuar borrador"],
            index=1 if latest_draft else 0,
            horizontal=True,
        )
    with c3:
        st.caption("Tip: Draft permite retomar. Final habilita export y se guarda en Drive.")

    if "active_report_id" not in st.session_state:
        st.session_state["active_report_id"] = None

    if action == "Nuevo informe":
        if st.button("➕ Crear informe"):
            rid = create_report(community_id, rep_date_str, user["email"])
            st.session_state["active_report_id"] = rid
            st.success("Informe creado ✅")
            st.rerun()
    else:
        if latest_draft and st.button("▶️ Abrir borrador"):
            st.session_state["active_report_id"] = latest_draft
            st.rerun()
        elif not latest_draft:
            st.info("No hay borrador reciente para esta comunidad.")

    report_id = st.session_state.get("active_report_id")
    if not report_id:
        st.stop()

    # Carga matriz checklist
    matrix = build_report_matrix(community_id)
    if matrix.empty:
        st.warning("Esta comunidad no tiene instalaciones/tareas. Admin debe configurarlas en Módulo Comunidades.")
        st.stop()

    # Estado actual del reporte
    rdf = get_reports_df()
    rrow = rdf[rdf["report_id"].astype(str) == str(report_id)]
    if rrow.empty:
        st.error("Informe no encontrado.")
        st.stop()
    rrow = rrow.iloc[0].to_dict()
    rstatus = str(rrow.get("status", "draft")).lower()

    # ReportItems existentes
    items_df = load_reportitems(report_id)
    # Index rápido por (category, installation, task)
    idx = {}
    if not items_df.empty:
        for _, rr in items_df.iterrows():
            key = (str(rr["category"]), str(rr["installation"]), str(rr["task"]))
            idx[key] = rr.to_dict()

    # UI: checklist 3 columnas visual
    st.markdown("---")
    st.markdown(f"### Informe: **{community_name}**  •  Estado: **{rstatus.upper()}**  •  ID: `{report_id}`")

    # Stats
    def _count_stats(df_items: pd.DataFrame):
        if df_items.empty:
            return 0, 0, len(matrix), len(matrix)
        s = df_items["status"].astype(str).str.lower()
        ok = int((s == "ok").sum())
        fail = int((s == "fail").sum())
        pending = int((s != "ok").sum() - fail)  # pending/otros
        return ok, fail, pending, ok + fail + pending

    okc, failc, pendc, totc = _count_stats(items_df)
    header_card(okc, failc, pendc)

    st.markdown("#### Checklist (estructura visual 3 columnas)")
    st.caption("Las filas con FALLA se marcan con fondo rojo suave. Guardas cuando termines (botón abajo).")

    # edición en memoria
    edited_rows = []

    # agrupar por categoría e instalación
    for cat in matrix["category"].drop_duplicates().tolist():
        st.markdown(f"### {cat}")
        mcat = matrix[matrix["category"] == cat]
        for inst_name in mcat["installation"].drop_duplicates().tolist():
            minst = mcat[mcat["installation"] == inst_name]

            for _, row in minst.iterrows():
                task = row["task"]
                key = (cat, inst_name, task)
                existing = idx.get(key, {})
                cur_status = str(existing.get("status", "pending")).lower() or "pending"
                cur_note = str(existing.get("note", "") or "")
                cur_photo_b64 = existing.get("photo_b64", "")
                cur_photo_mime = existing.get("photo_mime", "")

                bg = "#fee2e2" if cur_status == "fail" else "#ffffff"
                box = st.container(border=True)
                with box:
                    st.markdown(
                        f"""
                        <div style="background:{bg}; padding:10px 12px; border-radius:12px;">
                          <div style="font-weight:900;">{inst_name}</div>
                          <div style="opacity:0.75; font-size:12px;">Tarea: {task}</div>
                        </div>
                        """,
                        unsafe_allow_html=True,
                    )

                    cL, cM, cR = st.columns([2.4, 2.4, 1.7], gap="medium")

                    with cL:
                        st.write("**Categoría / Instalación / Tarea**")
                        st.caption(f"{cat} • {inst_name} • {task}")

                    with cM:
                        st.write("**Estado + Observación**")
                        status = st.radio(
                            "Estado",
                            options=["pending", "ok", "fail"],
                            index=["pending", "ok", "fail"].index(cur_status) if cur_status in ["pending", "ok", "fail"] else 0,
                            format_func=lambda v: {"pending": "Pendiente", "ok": "OK", "fail": "Falla"}[v],
                            horizontal=True,
                            key=f"st_{report_id}_{hash(key)}",
                            label_visibility="collapsed",
                        )
                        note = st.text_input(
                            "Observación",
                            value=cur_note,
                            placeholder="Observación breve…",
                            key=f"nt_{report_id}_{hash(key)}",
                            label_visibility="collapsed",
                        )

                    with cR:
                        st.write("**Foto (opcional)**")
                        up = st.file_uploader(
                            "Subir",
                            type=["png", "jpg", "jpeg"],
                            key=f"ph_{report_id}_{hash(key)}",
                            label_visibility="collapsed",
                        )
                        photo_bytes = None
                        photo_mime = ""
                        if up is not None:
                            photo_bytes = up.getvalue()
                            photo_mime = up.type or "image/jpeg"
                            st.image(photo_bytes, use_container_width=True, caption="Vista previa")
                        elif cur_photo_b64:
                            # muestra miniatura si ya existe
                            try:
                                pb = base64.b64decode(cur_photo_b64)
                                st.image(pb, use_container_width=True, caption="Foto guardada")
                            except Exception:
                                st.caption("Foto guardada (no visualizable).")

                # construye registro report item
                out = {
                    "report_item_id": existing.get("report_item_id") or str(uuid.uuid4()),
                    "report_id": report_id,
                    "category": cat,
                    "installation": inst_name,
                    "task": task,
                    "status": status,
                    "note": note,
                    "photo_mime": photo_mime or cur_photo_mime,
                    "photo_b64": base64.b64encode(photo_bytes).decode("utf-8") if photo_bytes else (cur_photo_b64 or ""),
                    "updated_at": _now_iso(),
                }
                edited_rows.append(out)

    # Guardar
    st.markdown("---")
    st.markdown("### Estado del informe")
    st.caption("Draft: permite retomar después. Final: habilita exportación.")
    new_status = st.radio("Estado", options=["draft", "final"], index=0 if rstatus == "draft" else 1, horizontal=True)
    if st.button("💾 Guardar cambios", use_container_width=True):
        try:
            df_save = pd.DataFrame(edited_rows, columns=REPORTITEMS_HEADERS).fillna("")
            save_reportitems(report_id, df_save)
            set_report_status(report_id, new_status)
            st.success("Guardado ✅")
            st.rerun()
        except Exception as e:
            st.error(str(e))

    # Export (solo si final)
    if new_status == "final":
        st.markdown("### Descargar / Guardar en Drive")
        st.caption("Export visual 3 columnas (categoría/instalación/tarea | estado/obs | foto). Si Drive está conectado, se sube a Mi unidad.")

        # reconstruye rows para export (desde edited_rows, no desde sheet)
        export_rows = []
        for r in edited_rows:
            pb = None
            if r.get("photo_b64"):
                try:
                    pb = base64.b64decode(r["photo_b64"])
                except Exception:
                    pb = None
            export_rows.append({
                "category": r["category"],
                "installation": r["installation"],
                "task": r["task"],
                "status": r["status"],
                "note": r["note"],
                "photo_bytes": pb,
                "photo_mime": r.get("photo_mime", ""),
            })

        col1, col2 = st.columns([1, 1])

        with col1:
            if st.button("📄 Generar PDF", use_container_width=True):
                try:
                    pdf_bytes = export_pdf_visual(community_name, rep_date_str, export_rows)
                    st.download_button(
                        "⬇️ Descargar PDF",
                        data=pdf_bytes,
                        file_name=f"informe_{community_name}_{rep_date_str}.pdf".replace(" ", "_"),
                        mime="application/pdf",
                        use_container_width=True,
                    )
                    # subir a Drive
                    try:
                        if DRIVE_REPORTS_FOLDER_ID:
                            fid, link = drive_upload_bytes(
                                pdf_bytes,
                                filename=f"informe_{community_name}_{rep_date_str}.pdf".replace(" ", "_"),
                                mime_type="application/pdf",
                                parent_folder_id=DRIVE_REPORTS_FOLDER_ID,
                            )
                            _update_row_by_id(
                                TAB_REPORTS, REPORTS_HEADERS, "report_id", report_id,
                                {"drive_pdf_file_id": fid, "drive_pdf_link": link, "updated_at": _now_iso()}
                            )
                            st.success("PDF subido a Drive ✅")
                            if link:
                                st.link_button("Abrir PDF en Drive", link)
                    except Exception as e:
                        st.info(f"PDF generado. Drive: {e}")
                except Exception as e:
                    st.error(str(e))

        with col2:
            if st.button("📝 Generar Word (DOCX)", use_container_width=True):
                try:
                    docx_bytes = export_docx_visual(community_name, rep_date_str, export_rows)
                    st.download_button(
                        "⬇️ Descargar DOCX",
                        data=docx_bytes,
                        file_name=f"informe_{community_name}_{rep_date_str}.docx".replace(" ", "_"),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                    # subir a Drive
                    try:
                        if DRIVE_REPORTS_FOLDER_ID:
                            fid, link = drive_upload_bytes(
                                docx_bytes,
                                filename=f"informe_{community_name}_{rep_date_str}.docx".replace(" ", "_"),
                                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                parent_folder_id=DRIVE_REPORTS_FOLDER_ID,
                            )
                            _update_row_by_id(
                                TAB_REPORTS, REPORTS_HEADERS, "report_id", report_id,
                                {"drive_docx_file_id": fid, "drive_docx_link": link, "updated_at": _now_iso()}
                            )
                            st.success("DOCX subido a Drive ✅")
                            if link:
                                st.link_button("Abrir DOCX en Drive", link)
                    except Exception as e:
                        st.info(f"DOCX generado. Drive: {e}")
                except Exception as e:
                    st.error(str(e))
    else:
        st.info("Para descargar, cambia el estado a FINAL y guarda.")


def module_resumen(user: dict):
    st.subheader("📊 Módulo Resumen (historial)")
    st.caption("Tabla de informes: Comunidad | Fecha | Usuario | Descarga (PDF/DOCX). Filtra por comunidad, usuario y rango de fechas.")

    comms = get_communities_df()
    reports = get_reports_df()
    if reports.empty:
        st.info("No hay informes.")
        return

    # joins
    comm_map = {}
    if not comms.empty:
        for _, r in comms.iterrows():
            comm_map[str(r["community_id"])] = str(r["name"])

    rep = reports.copy()
    rep["community_name"] = rep["community_id"].astype(str).map(comm_map).fillna(rep["community_id"].astype(str))
    rep["report_date"] = rep["report_date"].astype(str)

    # filtros permisos
    if user["role"] != ROLE_ADMIN:
        allowed_ids = set([x.strip() for x in (user.get("communities_csv") or "").split(",") if x.strip()])
        rep = rep[rep["community_id"].astype(str).isin(allowed_ids)].copy()

    if rep.empty:
        st.warning("No tienes informes accesibles según tus comunidades asignadas.")
        return

    # filtros UI
    col1, col2, col3 = st.columns([1.2, 1.2, 1.6])
    with col1:
        comm_filter = st.multiselect(
            "Filtrar comunidad",
            options=sorted(rep["community_name"].unique().tolist()),
        )
    with col2:
        user_filter = st.multiselect(
            "Filtrar por usuario",
            options=sorted(rep["created_by"].astype(str).unique().tolist()),
        )
    with col3:
        # rango fechas simple
        dmin = rep["report_date"].min()
        dmax = rep["report_date"].max()
        try:
            d1 = date.fromisoformat(dmin)
            d2 = date.fromisoformat(dmax)
        except Exception:
            d1, d2 = date.today() - timedelta(days=30), date.today()
        rng = st.date_input("Rango fechas", value=(d1, d2))

    if comm_filter:
        rep = rep[rep["community_name"].isin(comm_filter)].copy()
    if user_filter:
        rep = rep[rep["created_by"].isin(user_filter)].copy()

    if isinstance(rng, tuple) and len(rng) == 2:
        a, b = rng
        rep = rep[(rep["report_date"] >= a.isoformat()) & (rep["report_date"] <= b.isoformat())].copy()

    rep = rep.sort_values(["report_date", "updated_at"], ascending=[False, False])

    # tabla visual + links
    view = rep[[
        "community_name", "report_date", "status", "created_by",
        "drive_pdf_link", "drive_docx_link", "updated_at"
    ]].copy()

    st.dataframe(view, use_container_width=True)

    st.caption("Si el informe tiene link Drive, puedes abrirlo desde aquí:")
    for _, r in view.head(15).iterrows():
        line = f"**{r['community_name']}** • {r['report_date']} • {r['created_by']} • {str(r['status']).upper()}"
        st.write(line)
        c1, c2 = st.columns([1, 1])
        with c1:
            if str(r.get("drive_pdf_link", "")).strip():
                st.link_button("Abrir PDF", r["drive_pdf_link"])
        with c2:
            if str(r.get("drive_docx_link", "")).strip():
                st.link_button("Abrir DOCX", r["drive_docx_link"])


# ============================================================
# MAIN
# ============================================================

def main():
    # auth
    if not current_user():
        login_screen()
        return

    user = current_user()

    # sidebar
    st.sidebar.title("Menú")
    st.sidebar.write(f"**{user['email']}**")
    st.sidebar.caption(f"Rol: {user.get('role_label', user.get('role'))} ({user.get('role')})")

    if st.sidebar.button("Cerrar sesión", use_container_width=True):
        _logout()

    # Drive OAuth UI (solo admin)
    drive_oauth_ui(user)

    base_role = _role_base(user.get("role"))

    modules = []
    if base_role == ROLE_ADMIN:
        modules = ["Usuarios", "Comunidades", "Informes", "Resumen"]
    elif base_role == ROLE_EDITOR:
        modules = ["Informes", "Resumen"]
    else:
        modules = ["Resumen"]

    mod = st.sidebar.radio("Módulo", options=modules)

    # content
    if mod == "Usuarios":
        module_users(user)
    elif mod == "Comunidades":
        module_communities(user)
    elif mod == "Informes":
        module_informes(user)
    elif mod == "Resumen":
        module_resumen(user)

    st.markdown(
        "<div style='opacity:0.6; font-size:12px; margin-top:18px;'>Control Comunidades • Sheets+Drive • Esquema Tasks/Reports/ReportItems • Testing</div>",
        unsafe_allow_html=True,
    )


if __name__ == "__main__":
    try:
        ensure_schema()
        main()
    except Exception as e:
        st.error("Error crítico inicializando la app.")
        st.exception(e)